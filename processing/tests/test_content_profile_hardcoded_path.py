import unittest

from docx import Document as create_document
from docx.document import Document as DocxDocument

from processing.src.services.extractors.citations_extractor import extract_citations
from processing.src.services.extractors.content_profile_extractor import extract_content_profile
from processing.src.services.extractors.toc_extractor import extract_toc


HEADERS = [
    "Level",
    "Name",
    "Required",
    "Definition",
    "Example",
    "Note",
    "TOC Requirements",
    "SME Comments",
]


def _build_levels_doc(rows: list[tuple[str, str, str]]) -> DocxDocument:
    doc = create_document()
    table = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))

    for idx, header in enumerate(HEADERS):
        table.rows[0].cells[idx].text = header

    for row_index, (level, name, definition) in enumerate(rows, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = level
        cells[1].text = name
        cells[2].text = "True"
        cells[3].text = definition
        cells[4].text = name

    return doc


def _build_legacy_levels_doc() -> DocxDocument:
    doc = create_document()
    doc.add_paragraph("Structuring Requirements", style="Heading 1")
    doc.add_paragraph("Scope", style="Heading 2")
    doc.add_paragraph("Document Structure")
    doc.add_paragraph("Levels")

    lines = [
        "Level 0",
        "Required: True",
        "Definition: Hard coded – /CL",
        "Level 1",
        "Required: True",
        'Definition: Hardcoded – "/CLCMFCompdeNormas"',
        "Level 2",
        "Required: True",
        "Definition: document title",
        "Example: Compendio de Normas Contables para Bancos",
        "Level 3",
        "Required: False",
        "Name: Chapter",
        'Definition: "Capítulo" + incrementing uppercase letter',
        "Example: Capítulo A-1",
    ]
    for text in lines:
        doc.add_paragraph(text, style="List Paragraph" if ":" in text else None)

    doc.add_paragraph("Annotated Header Text Levels", style="Heading 3")
    doc.add_paragraph("The following Levels should have header text annotations captured: Level 2")
    doc.add_paragraph("Metadata", style="Heading 2")
    return doc


def _build_combined_citation_doc() -> DocxDocument:
    doc = create_document()
    table = doc.add_table(rows=4, cols=5)
    headers = ["Level", "Is Level Citable?", "Rules", "Source of Law Level (select one)", "SME Comments"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    rows = [
        ("1", "N", "", "", ""),
        ("2", "Y", "<Level 2> Example: Canadian Radio-television and Telecommunications Commission Unsolicited Telecommunications Rules", "X", ""),
        ("3", "Y", "<Level 2> + \" | \" + <Level 3> Example: Canadian Radio-television and Telecommunications Commission Unsolicited Telecommunications Rules | Part II", "", ""),
    ]
    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            table.rows[row_index].cells[col_index].text = value

    return doc


class HardcodedPathExtractionTests(unittest.TestCase):
    def test_extract_citations_supports_combined_citable_rules_table(self):
        result = extract_citations(_build_combined_citation_doc())
        refs = {row["level"]: row for row in result["references"]}

        self.assertEqual(refs["2"]["isCitable"], "Y")
        self.assertEqual(refs["2"]["sourceOfLaw"], "X")
        self.assertIn("<Level 2>", refs["2"]["citationRules"])
        self.assertIn("Part II", refs["3"]["citationRules"])

    def test_extract_toc_supports_plain_paragraph_levels_section(self):
        toc = extract_toc(_build_legacy_levels_doc())
        sections = toc["sections"]

        self.assertGreaterEqual(len(sections), 4)
        self.assertEqual([section["level"] for section in sections[:4]], ["0", "1", "2", "3"])
        self.assertEqual(sections[0]["path"], "/CL")
        self.assertEqual(sections[1]["path"], "/CLCMFCompdeNormas")

    def test_extract_toc_keeps_level_labels_and_paths(self):
        doc = _build_levels_doc(
            [
                ("Level 0", "CL", 'Hardcoded – "/CL"'),
                ("Level 1", "CMFCompdeNormas", 'Hardcoded – "/CMFCompdeNormas"'),
                ("Level 2", "Book", "Book number"),
            ]
        )

        result = extract_toc(doc)
        sections = result["sections"]

        self.assertEqual([section["level"] for section in sections[:3]], ["0", "1", "2"])
        self.assertEqual(sections[0]["path"], "/CL")
        self.assertEqual(sections[1]["path"], "/CMFCompdeNormas")

    def test_extract_content_profile_derives_hardcoded_path_from_level_zero_and_one(self):
        doc = _build_levels_doc(
            [
                ("Level 0", "CL", 'Hardcoded – "/CL"'),
                ("Level 1", "CMFCompdeNormas", 'Hardcoded – "/CMFCompdeNormas"'),
                ("Level 2", "Book", "Book number"),
            ]
        )

        result = extract_content_profile(doc)

        self.assertEqual(result["hardcoded_path"], "/CL/CMFCompdeNormas")

    def test_extract_content_profile_from_plain_levels_section(self):
        result = extract_content_profile(_build_legacy_levels_doc())

        self.assertEqual(result["hardcoded_path"], "/CL/CLCMFCompdeNormas")
        self.assertGreaterEqual(len(result["levels"]), 4)


if __name__ == "__main__":
    unittest.main()
