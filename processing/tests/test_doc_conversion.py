import base64
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.services import extractor
from src.services.extraction_diagnostics import build_extraction_diagnostics, build_format_fingerprint
from src.services.extractors.citations_extractor import extract_citations
from src.services.extractors.metadata_extractor import extract_metadata
from src.services.extractors.scope_extractor import extract_scope


MHTML_DOC_SAMPLE = """MIME-Version: 1.0
Content-Type: multipart/related; boundary=\"----=_Part_1\"

------=_Part_1
Content-Type: text/html; charset=UTF-8
Content-Transfer-Encoding: quoted-printable

<html>
  <body>
    <h1>Structuring Requirements</h1>
    <h2>Scope</h2>
    <p>The scope of regulatory content should be extracted.</p>
    <table>
      <tr>
        <td><strong>Document title</strong></td>
        <td><strong>Reference URL</strong></td>
      </tr>
      <tr>
        <td>Title 26: Internal Revenue</td>
        <td>https://www.ecfr.gov/</td>
      </tr>
    </table>
    <h1>Citation Format Requirements</h1>
    <table>
      <tr>
        <td><strong>Level</strong></td>
        <td><strong>Citation rules</strong></td>
        <td><strong>Source Of Law</strong></td>
        <td><strong>SME Comments</strong></td>
      </tr>
      <tr>
        <td>3</td>
        <td>Title 26 | Chapter I</td>
        <td>Chapter</td>
        <td>Looks good</td>
      </tr>
    </table>
    <h1>Metadata</h1>
    <table>
      <tr>
        <td><strong>Metadata Element</strong></td>
        <td><strong>Document Location</strong></td>
        <td><strong>SME Comments</strong></td>
      </tr>
      <tr>
        <td>Source Name</td>
        <td>Code of Federal Regulations</td>
        <td>Correct</td>
      </tr>
    </table>
  </body>
</html>
------=_Part_1--
"""


class ConvertDocToDocxTests(unittest.TestCase):
    def test_prefers_native_converter_before_mhtml_fallback(self):
        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(MHTML_DOC_SAMPLE)
            src_path = handle.name

        converted_path = None
        original_is_mhtml = extractor._is_mhtml_doc
        original_word = extractor._convert_doc_to_docx_with_word
        original_soffice = extractor._convert_doc_to_docx_with_soffice
        original_mhtml = extractor._convert_mhtml_doc_to_docx
        mhtml_called = False

        def fake_word(_src: str, dst: str) -> bool:
            doc = Document()
            scope = doc.add_table(rows=2, cols=2)
            scope.rows[0].cells[0].text = "Document title"
            scope.rows[0].cells[1].text = "Reference URL"
            scope.rows[1].cells[0].text = "Title 26: Internal Revenue"
            scope.rows[1].cells[1].text = "https://www.ecfr.gov/"

            metadata = doc.add_table(rows=2, cols=3)
            metadata.rows[0].cells[0].text = "Metadata Element"
            metadata.rows[0].cells[1].text = "Document Location"
            metadata.rows[0].cells[2].text = "SME Comments"
            metadata.rows[1].cells[0].text = "Source Name"
            metadata.rows[1].cells[1].text = "Code of Federal Regulations"
            metadata.rows[1].cells[2].text = "Correct"

            citation = doc.add_table(rows=2, cols=4)
            citation.rows[0].cells[0].text = "Level"
            citation.rows[0].cells[1].text = "Citation rules"
            citation.rows[0].cells[2].text = "Source Of Law"
            citation.rows[0].cells[3].text = "SME Comments"
            citation.rows[1].cells[0].text = "3"
            citation.rows[1].cells[1].text = "Title 26 | Chapter I"
            citation.rows[1].cells[2].text = "Chapter"
            citation.rows[1].cells[3].text = "Looks good"

            doc.save(dst)
            return True

        def fake_mhtml(_src: str, _dst: str) -> bool:
            nonlocal mhtml_called
            mhtml_called = True
            return False

        try:
            extractor._is_mhtml_doc = lambda _path: True
            extractor._convert_doc_to_docx_with_word = fake_word
            extractor._convert_doc_to_docx_with_soffice = lambda *_args, **_kwargs: False
            extractor._convert_mhtml_doc_to_docx = fake_mhtml

            converted_path = extractor.convert_doc_to_docx(src_path)

            self.assertIsNotNone(converted_path)
            self.assertFalse(mhtml_called, "HTML fallback should not run when native conversion succeeds")
            assert converted_path is not None
            doc = Document(converted_path)
            all_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
            self.assertIn("Code of Federal Regulations", all_cells)
        finally:
            extractor._is_mhtml_doc = original_is_mhtml
            extractor._convert_doc_to_docx_with_word = original_word
            extractor._convert_doc_to_docx_with_soffice = original_soffice
            extractor._convert_mhtml_doc_to_docx = original_mhtml
            Path(src_path).unlink(missing_ok=True)
            if converted_path:
                Path(converted_path).unlink(missing_ok=True)

    def test_mhtml_doc_converts_without_external_tools(self):
        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(MHTML_DOC_SAMPLE)
            src_path = handle.name

        converted_path = None
        original_word = extractor._convert_doc_to_docx_with_word
        original_soffice = extractor._convert_doc_to_docx_with_soffice

        try:
            extractor._convert_doc_to_docx_with_word = lambda *_args, **_kwargs: False
            extractor._convert_doc_to_docx_with_soffice = lambda *_args, **_kwargs: False

            converted_path = extractor.convert_doc_to_docx(src_path)

            self.assertIsNotNone(converted_path)
            assert converted_path is not None
            self.assertTrue(os.path.exists(converted_path))

            doc = Document(converted_path)
            paragraphs = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Structuring Requirements", paragraphs)
            self.assertIn("The scope of regulatory content should be extracted.", paragraphs)

            self.assertGreaterEqual(len(doc.tables), 3)
            all_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
            self.assertIn("Document title", all_cells)
            self.assertIn("Title 26: Internal Revenue", all_cells)
            self.assertIn("Citation rules", all_cells)
            self.assertIn("Metadata Element", all_cells)
            self.assertIn("Code of Federal Regulations", all_cells)
        finally:
            extractor._convert_doc_to_docx_with_word = original_word
            extractor._convert_doc_to_docx_with_soffice = original_soffice
            Path(src_path).unlink(missing_ok=True)
            if converted_path:
                Path(converted_path).unlink(missing_ok=True)

    def test_mhtml_doc_with_base64_unicode_charset_converts_without_external_tools(self):
        html_payload = """
<html>
  <body>
    <h1>State Administration of Foreign Exchange (CN.SAFE) Rules</h1>
    <table>
      <tr><td>Document title</td><td>Reference URL</td></tr>
      <tr><td>SAFE Rules</td><td>https://www.safe.gov.cn/</td></tr>
    </table>
    <table>
      <tr><td>Metadata Element</td><td>Document Location</td><td>SME Comments</td></tr>
      <tr><td>Source Name</td><td>State Administration of Foreign Exchange (国家外汇管理局)</td><td>Verified</td></tr>
    </table>
  </body>
</html>
""".strip()
        encoded_html = base64.b64encode(html_payload.encode("utf-16le")).decode("ascii")
        sample = f"""MIME-Version: 1.0
Content-Type: multipart/related; boundary=\"----=_Part_Unicode\"

------=_Part_Unicode
Content-Type: text/html; charset=\"unicode\"
Content-Transfer-Encoding: base64

{encoded_html}
------=_Part_Unicode--
"""

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(sample)
            src_path = handle.name

        converted_path = None
        original_word = extractor._convert_doc_to_docx_with_word
        original_soffice = extractor._convert_doc_to_docx_with_soffice

        try:
            extractor._convert_doc_to_docx_with_word = lambda *_args, **_kwargs: False
            extractor._convert_doc_to_docx_with_soffice = lambda *_args, **_kwargs: False

            converted_path = extractor.convert_doc_to_docx(src_path)

            self.assertIsNotNone(converted_path)
            assert converted_path is not None
            doc = Document(converted_path)
            all_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
            self.assertIn("SAFE Rules", all_cells)
            self.assertIn("State Administration of Foreign Exchange (国家外汇管理局)", all_cells)
        finally:
            extractor._convert_doc_to_docx_with_word = original_word
            extractor._convert_doc_to_docx_with_soffice = original_soffice
            Path(src_path).unlink(missing_ok=True)
            if converted_path:
                Path(converted_path).unlink(missing_ok=True)

    def test_mhtml_doc_preserves_inline_rich_formatting_in_citation_rules(self):
        sample = """MIME-Version: 1.0
Content-Type: multipart/related; boundary=\"----=_Part_Styled\"

------=_Part_Styled
Content-Type: text/html; charset=UTF-8
Content-Transfer-Encoding: quoted-printable

<html>
  <body>
    <h1>Citation Format Requirements</h1>
    <h2>Citation Standardization Rules</h2>
    <table>
      <tr>
        <td><strong>Level</strong></td>
        <td><strong>Citation rules</strong></td>
        <td><strong>Source Of Law</strong></td>
        <td><strong>SME Comments</strong></td>
      </tr>
      <tr>
        <td>2</td>
        <td><span style=\"color:#1d7afc;font-weight:bold\">Keep this text</span><br><span style=\"text-decoration:line-through\">Remove this old pattern</span></td>
        <td><span style=\"color:#ae2e24\">Level 2</span></td>
        <td><span style=\"text-decoration:line-through\">Deprecated note</span></td>
      </tr>
    </table>
  </body>
</html>
------=_Part_Styled--
"""

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(sample)
            src_path = handle.name

        converted_path = None
        original_word = extractor._convert_doc_to_docx_with_word
        original_soffice = extractor._convert_doc_to_docx_with_soffice

        try:
            extractor._convert_doc_to_docx_with_word = lambda *_args, **_kwargs: False
            extractor._convert_doc_to_docx_with_soffice = lambda *_args, **_kwargs: False

            converted_path = extractor.convert_doc_to_docx(src_path)
            self.assertIsNotNone(converted_path)
            assert converted_path is not None

            extracted = extractor.extract_all(converted_path)
            refs = extracted.get("citations", {}).get("references", [])
            level_two = next((row for row in refs if row.get("level") == "2"), None)

            self.assertIsNotNone(level_two)
            assert level_two is not None
            self.assertIn("<strong>", level_two["citationRules"])
            self.assertRegex(level_two["citationRules"], r"(?i)color\s*:\s*#1d7afc")
            self.assertIn("<s>", level_two["citationRules"])
            self.assertIn("<s>", level_two["smeComments"])
        finally:
            extractor._convert_doc_to_docx_with_word = original_word
            extractor._convert_doc_to_docx_with_soffice = original_soffice
            Path(src_path).unlink(missing_ok=True)
            if converted_path:
                Path(converted_path).unlink(missing_ok=True)

    def test_scope_extraction_preserves_content_url_note_text(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=6)
        headers = ["Document Title", "Reference URL", "Content URL", "Issuing Authority", "ASRB ID", "SME Comments"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        row = table.rows[1].cells
        row[0].text = "SAFE Rule"
        row[1].text = "https://www.safe.gov.cn/reference"
        row[2].text = (
            "https://www.safe.gov.cn/source\n\n"
            "The rule is only found in the attachment sub-link\n"
            "1. 外债登记管理办法\n\n"
            "https://www.safe.gov.cn/file.pdf"
        )
        row[3].text = "State Administration of Foreign Exchange (SAFE) / China"
        row[4].text = "ASRB279"
        row[5].text = ""

        extracted = extract_scope(doc)
        self.assertEqual(len(extracted["in_scope"]), 1)
        entry = extracted["in_scope"][0]

        self.assertEqual(entry["content_url"], "https://www.safe.gov.cn/file.pdf")
        self.assertEqual(
            entry.get("content_note"),
            "The rule is only found in the attachment sub-link\n1. 外债登记管理办法",
        )

    def test_scope_extraction_moves_asrb_ids_out_of_sme_comments(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=6)
        headers = ["Document Title", "Reference URL", "Content URL", "Issuing Authority", "ASRB ID", "SME Comments"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        row = table.rows[1].cells
        row[0].text = "SAFE Rule"
        row[1].text = "https://www.safe.gov.cn/reference"
        row[2].text = "https://www.safe.gov.cn/file.pdf"
        row[3].text = "State Administration of Foreign Exchange (SAFE) / China"
        row[4].text = ""
        row[5].text = "ASRB279; validate attachment coverage"

        extracted = extract_scope(doc)
        self.assertEqual(len(extracted["in_scope"]), 1)
        entry = extracted["in_scope"][0]

        self.assertEqual(entry["asrb_id"], "ASRB279")
        self.assertEqual(entry["sme_comments"], "validate attachment coverage")

    def test_citation_extraction_ignores_citable_table_notes_when_rules_table_exists(self):
        doc = Document()

        citable = doc.add_table(rows=3, cols=3)
        citable.rows[0].cells[0].text = "Level"
        citable.rows[0].cells[1].text = "Is Level Citable?"
        citable.rows[0].cells[2].text = "SME Comments - if citation rules need changes, specify here"
        citable.rows[1].cells[0].text = "2"
        citable.rows[1].cells[1].text = "Y"
        citable.rows[1].cells[2].text = "Looks good"
        citable.rows[2].cells[0].text = "3"
        citable.rows[2].cells[1].text = "N"
        citable.rows[2].cells[2].text = "No citations"

        rules = doc.add_table(rows=3, cols=4)
        rules.rows[0].cells[0].text = "Level"
        rules.rows[0].cells[1].text = "Citation Rules"
        rules.rows[0].cells[2].text = "Source Of Law"
        rules.rows[0].cells[3].text = "SME Comments"
        rules.rows[1].cells[0].text = "2"
        rules.rows[1].cells[1].text = "Level 2 | Level 3"
        rules.rows[1].cells[2].text = "Level 2"
        rules.rows[1].cells[3].text = "Approved"
        rules.rows[2].cells[0].text = "3"
        rules.rows[2].cells[1].text = "Level 2 | Level 3 | Level 4"
        rules.rows[2].cells[2].text = "Level 2"
        rules.rows[2].cells[3].text = "Use full chain"

        extracted = extract_citations(doc)
        refs = {row["level"]: row for row in extracted["references"]}

        self.assertEqual(refs["2"]["isCitable"], "Y")
        self.assertEqual(refs["3"]["isCitable"], "N")
        self.assertEqual(refs["2"]["citationRules"], "Level 2 | Level 3")
        self.assertEqual(refs["2"]["sourceOfLaw"], "Level 2")

    def test_metadata_extraction_ignores_non_metadata_tables(self):
        doc = Document()
        stray = doc.add_table(rows=2, cols=2)
        stray.rows[0].cells[0].text = "Document Title"
        stray.rows[0].cells[1].text = "Reference URL"
        stray.rows[1].cells[0].text = "Content Type"
        stray.rows[1].cells[1].text = "Bogus scope value"

        metadata_table = doc.add_table(rows=3, cols=3)
        metadata_table.rows[0].cells[0].text = "Metadata Element"
        metadata_table.rows[0].cells[1].text = "Document Location"
        metadata_table.rows[0].cells[2].text = "SME Comments"
        metadata_table.rows[1].cells[0].text = "Source Name"
        metadata_table.rows[1].cells[1].text = "Real Source"
        metadata_table.rows[2].cells[0].text = "Content Type"
        metadata_table.rows[2].cells[1].text = "Regulation"

        extracted = extract_metadata(doc)

        self.assertEqual(extracted["content_category_name"], "Real Source")
        self.assertEqual(extracted["content_type"], "Regulation")

    def test_build_diagnostics_and_format_fingerprint(self):
        extracted = {
            "format": "new",
            "scope": {"in_scope": [], "out_of_scope": [{"document_title": "Archived row", "strikethrough": True}]},
            "toc": {"sections": [], "citationStyleGuide": {"rows": []}},
            "citations": {"references": [{"level": "2", "citationRules": "", "sourceOfLaw": "", "isCitable": "", "smeComments": ""}]},
        }

        diagnostics = build_extraction_diagnostics(extracted, image_count=0)
        fingerprint = build_format_fingerprint(filename="sample.doc", original_suffix=".doc", detected_format="new", is_mhtml_doc=True)

        self.assertTrue(any(item["code"] == "scope_all_rows_excluded" for item in diagnostics["warnings"]))
        self.assertTrue(any(item["code"] == "citations_incomplete" for item in diagnostics["warnings"]))
        self.assertEqual(fingerprint["container"], "mhtml-doc")
        self.assertEqual(fingerprint["template"], "new")

    def test_build_diagnostics_flags_partial_citations_missing_metadata_and_scope_review(self):
        extracted = {
            "format": "new",
            "metadata": {
                "content_category_name": "Manual V2",
                "issuing_agency": "",
                "content_uri": "",
            },
            "scope": {
                "in_scope": [
                    {"document_title": "", "regulator_url": "", "content_url": "https://example.com/rule"}
                ],
                "out_of_scope": [
                    {"document_title": "Archived rule", "strikethrough": True},
                    {"document_title": "Retired rule", "strikethrough": True},
                ],
            },
            "toc": {"sections": [{"section": "Overview"}], "citationStyleGuide": {"rows": [{"label": "Use title and chapter"}] }},
            "citations": {
                "references": [
                    {"level": "2", "citationRules": "Title | Chapter", "sourceOfLaw": "Act", "isCitable": "Y", "smeComments": ""},
                    {"level": "3", "citationRules": "", "sourceOfLaw": "", "isCitable": "", "smeComments": ""},
                ]
            },
        }

        diagnostics = build_extraction_diagnostics(extracted, image_count=1)
        codes = {item["code"] for item in diagnostics["warnings"]}

        self.assertIn("citations_partial", codes)
        self.assertIn("metadata_key_fields_missing", codes)
        self.assertIn("scope_review_recommended", codes)


if __name__ == "__main__":
    unittest.main()
