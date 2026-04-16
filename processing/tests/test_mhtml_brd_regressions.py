import tempfile
import textwrap
import unittest
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from src.services.extractor import _fallback_from_text, convert_doc_to_docx, extract_all
from src.services.extractors.scope_extractor import _is_non_data_scope_row, extract_scope, extract_scope_from_file
from src.services.extractors.image_extractor import extract_and_store_images_from_mhtml
from src.services.extractors.toc_extractor import extract_toc
from src.services.extractors.citations_extractor import _normalise_citation_rule, extract_citations
from src.services.extractors.base import extract_url_and_note_from_text


_TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
    "w8AAusB9Y9l9n8AAAAASUVORK5CYII="
)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class MhtmlBrdRegressionTests(unittest.TestCase):
    def test_raw_text_fallback_handles_noisy_toc_heading_variants(self):
        raw_text = textwrap.dedent(
            """\
            {}ToC* - Sorting order
            SME Checkpoint
            Sort numerically in descending order.

            *ToC - Hiding levels (Tech Only)
            Note:
            Level 8-14 not to be included in the TOC.
            """
        )

        extracted = _fallback_from_text(raw_text, "new")
        toc = extracted.get("toc", {})

        self.assertIn("tocSortingOrder", toc)
        self.assertIn("Sort numerically in descending order.", toc["tocSortingOrder"])
        self.assertIn("tocHidingLevels", toc)
        self.assertIn("Level 8-14 not to be included in the TOC.", toc["tocHidingLevels"])

    def test_raw_text_fallback_preserves_citation_and_toc_context_fields(self):
        raw_text = textwrap.dedent(
            """\
            Table of Contents
            1 Scope
            2 Document Structure

            Citation Style Guide Link
            SME Checkpoint
            Product Owner | Raut, Divya
            Source URL | https://example.com/style-guide

            ToC - Sorting Order
            Sort numerically in descending order.

            ToC - Hiding levels (Tech Only)
            Level 8-14 not to be included in the TOC.
            """
        )

        extracted = _fallback_from_text(raw_text, "new")
        toc = extracted.get("toc", {})

        self.assertIn("citationStyleGuide", toc)
        self.assertEqual(toc["citationStyleGuide"]["description"], "SME Checkpoint")
        self.assertEqual(toc["citationStyleGuide"]["rows"][0]["label"], "Product Owner")
        self.assertEqual(toc["citationStyleGuide"]["rows"][0]["value"], "Raut, Divya")
        self.assertEqual(toc["tocSortingOrder"], "Sort numerically in descending order.")
        self.assertEqual(toc["tocHidingLevels"], "Level 8-14 not to be included in the TOC.")

    def test_scope_rows_with_real_titles_are_not_dropped_by_generic_sme_comments(self):
        self.assertFalse(
            _is_non_data_scope_row(
                "Manual de Normas B3",
                "https://www.b3.com.br/pt_br/regulacao/estrutura-normativa/regulamentos-e-manuais/negociacao.htm",
                "https://www.b3.com.br/data/files/example.pdf",
                'Please, rename the document title for the following: "Manual de Normas Plataforma de Negociação do Balcão B3"',
            )
        )
        self.assertFalse(
            _is_non_data_scope_row(
                "Manual de Certificação de Profissionais B3",
                "https://www.b3.com.br/pt_br/b3/educacao/certificacoes-pqo/sobre-a-certificacao/",
                "https://www.b3.com.br/data/files/example-2.pdf",
                "The following link is not working: https://example.invalid/file.pdf",
            )
        )

    def test_scope_preserves_struck_rows_before_first_url_and_keeps_sme_comments(self):
        doc = Document()
        doc.add_heading("Scope", level=1)
        table = doc.add_table(rows=3, cols=6)
        headers = [
            "Document Title", "Reference URL", "Content URL", "Issuing Authority", "ASRB ID", "SME Comments"
        ]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        struck_title = table.rows[1].cells[0].paragraphs[0]
        struck_title.clear()
        run = struck_title.add_run("Withdrawn German Parliament source")
        run.font.strike = True
        table.rows[1].cells[5].text = "SME requested removal from monitoring"

        row = table.rows[2].cells
        row[0].text = "Active German Parliament source"
        row[1].text = "https://example.com/reference"
        row[2].text = "https://example.com/content.pdf"
        row[5].text = "Still monitored"

        scope = extract_scope(doc)

        self.assertEqual(len(scope["out_of_scope"]), 1)
        self.assertIn("Withdrawn German Parliament source", scope["out_of_scope"][0]["document_title"])
        self.assertIn("<s>", scope["out_of_scope"][0]["document_title"])
        self.assertEqual(scope["out_of_scope"][0]["sme_comments"], "SME requested removal from monitoring")
        self.assertEqual(len(scope["in_scope"]), 1)

    def test_citation_style_guide_preserves_hyperlinked_contributor_names_with_dates(self):
        doc = Document()
        doc.add_heading("Citation Style Guide Link", level=1)
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Contributors"

        para = table.rows[0].cells[1].paragraphs[0]
        para.clear()
        _add_hyperlink(para, "Raut, Divya", "https://example.com/divya")
        para.add_run(" (Nov 13, 2025), ")
        _add_hyperlink(para, "Haibach, Julia", "https://example.com/julia")
        para.add_run(" (Nov 17, 2025)")

        table.rows[1].cells[0].text = "Product Owner"
        table.rows[1].cells[1].text = "Example Owner"

        toc = extract_toc(doc)
        citation_guide = toc.get("citationStyleGuide") or {}
        rows = citation_guide.get("rows") or []

        self.assertEqual(rows[0]["label"], "Contributors")
        self.assertIn("Raut, Divya", rows[0]["value"])
        self.assertIn("Haibach, Julia", rows[0]["value"])
        self.assertIn("Nov 13, 2025", rows[0]["value"])
        self.assertIn("Nov 17, 2025", rows[0]["value"])

    def test_mhtml_scope_preserves_red_rich_text_rows(self):
        mhtml = textwrap.dedent(
            '''\
            MIME-Version: 1.0
            Content-Type: multipart/related; boundary="BOUNDARY"

            --BOUNDARY
            Content-Type: text/html; charset=UTF-8
            Content-Transfer-Encoding: quoted-printable
            Content-Location: file:///C:/exported.html

            <html>
              <body>
                <h2>Scope</h2>
                <table>
                  <tr>
                    <th>Document Title</th>
                    <th>Reference URL</th>
                    <th>Content URL</th>
                    <th>Issuing Authority</th>
                    <th>ASRB ID</th>
                    <th>SME Comments</th>
                    <th>Initial / Evergreen</th>
                    <th>Date of Ingestion</th>
                  </tr>
                  <tr>
                    <td><p><span style="color:#ae2e24;"><strong>Ordonnance test</strong></span></p></td>
                    <td><p><span style="color:#ae2e24;"><a href="https://example.com/reference">https://example.com/reference</a></span></p></td>
                    <td><p><span style="color:#ae2e24;"><a href="https://example.com/content">https://example.com/content</a></span></p></td>
                    <td><p>Swiss Authority</p></td>
                    <td><p>ASRB-321</p></td>
                    <td><p><span style="color:#ae2e24;"><strong>Needs review</strong></span></p></td>
                    <td><p><span style="color:#ae2e24;"><strong>Evergreen</strong></span></p></td>
                    <td><p><span style="color:#ae2e24;"><strong>WIP</strong></span></p></td>
                  </tr>
                </table>
              </body>
            </html>

            --BOUNDARY--
            '''
        )

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(mhtml)
            doc_path = handle.name

        try:
            scope = extract_scope_from_file(doc_path)
        finally:
            Path(doc_path).unlink(missing_ok=True)

        self.assertEqual(len(scope["in_scope"]), 1)
        row = scope["in_scope"][0]
        self.assertIn("color:#ae2e24", row["document_title"])
        self.assertIn("https://example.com/reference", row["regulator_url"])
        self.assertIn("color:#ae2e24", row["initial_evergreen"])
        self.assertIn("color:#ae2e24", row["date_of_ingestion"])

    def test_extracts_metadata_images_from_mhtml_confluence_exports(self):
        mhtml = textwrap.dedent(
            f'''\
            MIME-Version: 1.0
            Content-Type: multipart/related; boundary="BOUNDARY"

            --BOUNDARY
            Content-Type: text/html; charset=UTF-8
            Content-Transfer-Encoding: quoted-printable
            Content-Location: file:///C:/exported.html

            <html>
              <body>
                <h1>Metadata</h1>
                <table>
                  <tr><th>Metadata Element</th><th>Document Location</th></tr>
                  <tr>
                    <td><p><strong>Publication Date</strong></p></td>
                    <td>
                      <p>
                        Publication date can usually be found on the first page.<br>
                        <img src="pubdate-image.png" width="32" height="16" />
                      </p>
                    </td>
                  </tr>
                </table>
              </body>
            </html>

            --BOUNDARY
            Content-Type: image/png
            Content-Transfer-Encoding: base64
            Content-Location: file:///C:/pubdate-image.png

            {_TINY_PNG_BASE64}
            --BOUNDARY--
            '''
        )

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(mhtml)
            doc_path = handle.name

        try:
            images = extract_and_store_images_from_mhtml(doc_path, brd_id="TEST-MHTML")
        finally:
            Path(doc_path).unlink(missing_ok=True)

        self.assertEqual(len(images), 1)
        self.assertEqual(images[0]["section"], "metadata")
        self.assertEqual(images[0]["fieldLabel"], "Publication Date")
        self.assertEqual(images[0]["mimeType"], "image/png")

    def test_octet_stream_mhtml_images_are_upgraded_to_png_for_display(self):
        mhtml = textwrap.dedent(
            f'''\
            MIME-Version: 1.0
            Content-Type: multipart/related; boundary="BOUNDARY"

            --BOUNDARY
            Content-Type: text/html; charset=UTF-8
            Content-Transfer-Encoding: quoted-printable
            Content-Location: file:///C:/exported.html

            <html>
              <body>
                <h1>Metadata</h1>
                <table>
                  <tr><th>Metadata Element</th><th>Document Location</th></tr>
                  <tr>
                    <td><p><strong>Publication Date</strong></p></td>
                    <td>
                      <p>
                        <img
                          src="737319a2dabd45e08220ff2e024f3246364948f82d0760e8336b5b3610779671"
                          data-image-src="/confluence/download/attachments/1/pubdate-image.png?version=1"
                          data-linked-resource-content-type="image/png"
                        />
                      </p>
                    </td>
                  </tr>
                </table>
              </body>
            </html>

            --BOUNDARY
            Content-Type: application/octet-stream
            Content-Transfer-Encoding: base64
            Content-Location: file:///C:/737319a2dabd45e08220ff2e024f3246364948f82d0760e8336b5b3610779671

            {_TINY_PNG_BASE64}
            --BOUNDARY--
            '''
        )

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(mhtml)
            doc_path = handle.name

        try:
            images = extract_and_store_images_from_mhtml(doc_path, brd_id="TEST-MHTML-OCTET")
        finally:
            Path(doc_path).unlink(missing_ok=True)

        self.assertEqual(len(images), 1)
        self.assertEqual(images[0]["mimeType"], "image/png")
        self.assertTrue(images[0]["mediaName"].endswith(".png"))

    def test_citation_rule_soft_breaks_collapse_to_single_line_spacing(self):
        raw = 'This is the example for this level:\n\t<Level 2> + "," + <Level 3>\nCrystal-Based on discussion with Paula'
        self.assertEqual(
            _normalise_citation_rule(raw),
            'This is the example for this level: <Level 2> + "," + <Level 3> Crystal-Based on discussion with Paula',
        )

    def test_scope_url_with_parentheses_keeps_full_pdf_link(self):
        text = 'https://www.b3.com.br/data/files/82/91/0D/02/00D8C810719CE3C8DC0D8AA8/OC%20214-2023%20PRE%20PEC%20do%20PQO%20(PT).pdf'
        url, note = extract_url_and_note_from_text(text)
        self.assertEqual(url, text)
        self.assertEqual(note, '')

    def test_extract_all_preserves_metadata_content_uri_note_alongside_link(self):
        doc = Document()
        doc.add_heading("Metadata", level=1)
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Metadata Element"
        table.rows[0].cells[1].text = "Document Location"
        table.rows[0].cells[2].text = "SME Comments"
        table.rows[1].cells[0].text = "Content URI"
        table.rows[1].cells[1].text = (
            "URL of the specific Document (e.g.)\n"
            "https://www.b3.com.br/data/files/example.pdf"
        )
        table.rows[1].cells[2].text = "Ok"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            doc.save(handle.name)
            doc_path = handle.name

        try:
            extracted = extract_all(doc_path)
        finally:
            Path(doc_path).unlink(missing_ok=True)

        metadata = extracted["metadata"]
        self.assertEqual(metadata["Content URI"], "https://www.b3.com.br/data/files/example.pdf")
        self.assertEqual(metadata.get("Content URI Note"), "URL of the specific Document (e.g.)")

    def test_extract_all_preserves_scope_evergreen_and_ingestion_columns(self):
        doc = Document()
        doc.add_heading("Scope", level=1)
        table = doc.add_table(rows=2, cols=8)
        headers = [
            "Document Title", "Reference URL", "Content URL", "Issuing Authority",
            "ASRB ID", "SME Comments", "Initial / Evergreen", "Date of Ingestion",
        ]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        row = table.rows[1].cells
        row[0].text = "Banco Central Scope"
        row[1].text = "https://example.com/reference"
        row[2].text = "https://example.com/content.pdf"
        row[3].text = "BCRA"
        row[4].text = "ASRB-123"
        row[5].text = "Tracked for ingestion"
        row[6].text = "Evergreen"
        row[7].text = "2024-01-15"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            doc.save(handle.name)
            doc_path = handle.name

        try:
            extracted = extract_all(doc_path)
        finally:
            Path(doc_path).unlink(missing_ok=True)

        self.assertEqual(len(extracted["scope"]["in_scope"]), 1)
        scope_row = extracted["scope"]["in_scope"][0]
        self.assertEqual(scope_row["initial_evergreen"], "Evergreen")
        self.assertEqual(scope_row["date_of_ingestion"], "2024-01-15")

    def test_toc_requirements_preserve_italic_and_strikethrough_markup(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=8)
        headers = [
            "Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"
        ]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        row = table.rows[1].cells
        row[0].text = "8"
        row[1].text = "Paragraph"
        row[2].text = "True"
        row[3].text = "numeric"
        row[4].text = "8.1"
        row[5].text = "note"

        para = row[6].paragraphs[0]
        para.clear()
        italic_run = para.add_run("Visible till Level 5")
        italic_run.italic = True
        para.add_run(" ")
        strike_run = para.add_run("not to be included in the TOC")
        strike_run.font.strike = True

        toc = extract_toc(doc)
        self.assertEqual(len(toc["sections"]), 1)
        toc_requirements = toc["sections"][0]["tocRequirements"]
        self.assertIn("Visible till Level 5", toc_requirements)
        self.assertIn("not to be included in the TOC", toc_requirements)
        self.assertRegex(toc_requirements, r"<(em|i)>Visible till Level 5</(em|i)>")
        self.assertRegex(toc_requirements, r"<(s|strike|del)>not to be included in the TOC</(s|strike|del)>")

    def test_toc_cells_preserve_style_based_strong_and_emphasis_markup(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=8)
        headers = [
            "Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"
        ]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        row = table.rows[1].cells
        row[0].text = "4"
        row[1].text = "Section"
        row[2].text = "True"
        row[3].text = "definition"
        row[5].text = "note"
        row[6].text = "Only Level identifier to be captured"
        row[7].text = "Ok"

        para = row[4].paragraphs[0]
        para.clear()
        strong_run = para.add_run("Bold sample")
        strong_run.style = "Strong"
        para.add_run(" ")
        emphasis_run = para.add_run("Italic sample")
        emphasis_run.style = "Emphasis"

        toc = extract_toc(doc)
        self.assertEqual(len(toc["sections"]), 1)
        example = toc["sections"][0]["example"]
        self.assertIn("Bold sample", example)
        self.assertIn("Italic sample", example)
        self.assertRegex(example, r"<(strong|b)>Bold sample</(strong|b)>")
        self.assertRegex(example, r"<(em|i)>Italic sample</(em|i)>")

    def test_mhtml_doc_conversion_keeps_toc_rich_text_markup(self):
        mhtml = textwrap.dedent(
            '''\
            MIME-Version: 1.0
            Content-Type: multipart/related; boundary="BOUNDARY"

            --BOUNDARY
            Content-Type: text/html; charset=UTF-8
            Content-Transfer-Encoding: quoted-printable
            Content-Location: file:///C:/exported.html

            <html>
              <body>
                <h1>Document Structure</h1>
                <h2>Levels</h2>
                <table>
                  <tr>
                    <th>Level</th><th>Name</th><th>Required</th><th>Definition</th>
                    <th>Example</th><th>Note</th><th>TOC Requirements</th><th>SME Comments</th>
                  </tr>
                  <tr>
                    <td>8</td><td>Paragraph</td><td>True</td><td>numeric</td>
                    <td>8.1</td><td>Note</td>
                    <td><em>Visible till Level 5</em> <s>not to be included in the TOC</s></td>
                    <td>Comment</td>
                  </tr>
                </table>
              </body>
            </html>
            --BOUNDARY--
            '''
        )

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(mhtml)
            doc_path = handle.name

        converted_path = None
        try:
            converted_path = convert_doc_to_docx(doc_path)
            self.assertIsNotNone(converted_path)
            converted = Document(converted_path)
            toc = extract_toc(converted)
        finally:
            Path(doc_path).unlink(missing_ok=True)
            if converted_path:
                Path(converted_path).unlink(missing_ok=True)

        self.assertEqual(len(toc["sections"]), 1)
        toc_requirements = toc["sections"][0]["tocRequirements"]
        self.assertRegex(toc_requirements, r"<(em|i)>Visible till Level 5</(em|i)>")
        self.assertRegex(toc_requirements, r"<(s|strike|del)>not to be included in the TOC</(s|strike|del)>")

    def test_section_checkpoint_extractors_ignore_page_toc_entries(self):
        doc = Document()
        doc.add_paragraph("Scope")
        doc.add_paragraph("*ToC - Sorting order")
        doc.add_paragraph("Citable Levels")
        doc.add_paragraph("Citation Standardization Rules")
        doc.add_paragraph("Metadata")

        doc.add_heading("Scope", level=2)
        doc.add_paragraph("SME Check-point")
        scope_note = doc.add_paragraph()
        scope_note.add_run("SMEs to check if weblink is correct for document titles under the source name.")
        scope_table = doc.add_table(rows=2, cols=4)
        scope_table.rows[0].cells[0].text = "Document title"
        scope_table.rows[0].cells[1].text = "Reference URL"
        scope_table.rows[0].cells[2].text = "Content URL"
        scope_table.rows[0].cells[3].text = "SME Comments"
        scope_table.rows[1].cells[0].text = "Sample scope title"

        doc.add_heading("*ToC - Sorting order", level=2)
        doc.add_paragraph("Please sort alphabetically by Document Title.")

        doc.add_heading("Citable Levels", level=2)
        doc.add_paragraph("Please indicate which levels are citable.")
        citable_table = doc.add_table(rows=2, cols=3)
        citable_table.rows[0].cells[0].text = "Level"
        citable_table.rows[0].cells[1].text = "Is Level Citable?"
        citable_table.rows[0].cells[2].text = "SME Comments"
        citable_table.rows[1].cells[0].text = "2"
        citable_table.rows[1].cells[1].text = "Y"

        doc.add_heading("Citation Standardization Rules", level=2)
        doc.add_paragraph("Citation rules stand for how the citations should appear in ELA.")
        rules_table = doc.add_table(rows=2, cols=4)
        rules_table.rows[0].cells[0].text = "Level"
        rules_table.rows[0].cells[1].text = "Citation Rules"
        rules_table.rows[0].cells[2].text = "Source of Law"
        rules_table.rows[0].cells[3].text = "SME Comments"
        rules_table.rows[1].cells[0].text = "2"
        rules_table.rows[1].cells[1].text = "<Level 2>"

        scope = extract_scope(doc)
        toc = extract_toc(doc)
        citations = extract_citations(doc)

        self.assertIn("weblink is correct", scope.get("smeCheckpoint", ""))
        self.assertNotIn("Metadata", scope.get("smeCheckpoint", ""))
        self.assertIn("sort alphabetically", toc.get("tocSortingOrder", ""))
        self.assertNotIn("Citable Levels", toc.get("tocSortingOrder", ""))
        self.assertIn("which levels are citable", citations.get("citationLevelSmeCheckpoint", ""))
        self.assertIn("citations should appear in ELA", citations.get("citationRulesSmeCheckpoint", ""))

    def test_extract_toc_includes_citation_style_and_toc_context_blocks(self):
        doc = Document()
        doc.add_heading("Citation Style Guide Link", level=2)
        doc.add_paragraph("SME Checkpoint - validate owner and status details.")

        guide_table = doc.add_table(rows=2, cols=2)
        guide_table.rows[0].cells[0].text = "Product Owner"
        guide_table.rows[0].cells[1].text = "Raut, Divya"
        guide_table.rows[1].cells[0].text = "Status"
        guide_table.rows[1].cells[1].text = "DRAFT BRDS"

        doc.add_heading("*ToC - Sorting order", level=2)
        doc.add_paragraph("Communications A should appear before Communications B.")

        doc.add_heading("*ToC - Hiding levels (Tech Only)", level=2)
        doc.add_paragraph("Level 8-14 not to be included in the TOC.")

        doc.add_heading("Document Structure", level=1)
        levels = doc.add_table(rows=2, cols=8)
        headers = ["Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"]
        for idx, header in enumerate(headers):
            levels.rows[0].cells[idx].text = header
        row = levels.rows[1].cells
        row[0].text = "2"
        row[1].text = "Title"
        row[2].text = "True"
        row[3].text = "document title"
        row[4].text = "Sample"

        extracted = extract_toc(doc)
        self.assertEqual(extracted["citationStyleGuide"]["rows"][0]["label"], "Product Owner")
        self.assertEqual(extracted["citationStyleGuide"]["rows"][0]["value"], "Raut, Divya")
        self.assertIn("Communications A", extracted["tocSortingOrder"])
        self.assertIn("Level 8-14", extracted["tocHidingLevels"])
        self.assertEqual(extracted["sections"][0]["level"], "2")

    def test_extract_toc_preserves_citation_guide_rows_with_unstyled_headings(self):
        doc = Document()
        doc.add_paragraph("Citation Style Guide Link")

        guide_table = doc.add_table(rows=4, cols=2)
        guide_table.rows[0].cells[0].text = "Document Type"
        guide_table.rows[1].cells[0].text = "Contributors"
        guide_table.rows[2].cells[0].text = "Innodata Last Edit Date"
        guide_table.rows[3].cells[0].text = "Innodata Fields Changed"

        doc.add_paragraph("Document Structure")
        levels = doc.add_table(rows=2, cols=8)
        headers = ["Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"]
        for idx, header in enumerate(headers):
            levels.rows[0].cells[idx].text = header
        levels.rows[1].cells[0].text = "2"
        levels.rows[1].cells[1].text = "Title"
        levels.rows[1].cells[6].text = "sort alphabetically"

        extracted = extract_toc(doc)
        guide = extracted.get("citationStyleGuide") or {}
        rows = guide.get("rows") or []

        self.assertEqual([row["label"] for row in rows[:4]], [
            "Document Type",
            "Contributors",
            "Innodata Last Edit Date",
            "Innodata Fields Changed",
        ])
        self.assertTrue(all(row.get("value", "") == "" for row in rows[:4]))
        self.assertNotIn("Document Type", guide.get("description", ""))
        self.assertEqual(extracted["sections"][0]["level"], "2")

    def test_extract_toc_collects_citation_guide_rows_from_multiple_tables(self):
        doc = Document()
        doc.add_heading("Citation Style Guide Link", level=2)

        description_table = doc.add_table(rows=1, cols=2)
        description_table.rows[0].cells[0].text = "SME Checkpoint"
        desc_para = description_table.rows[0].cells[1].paragraphs[0]
        desc_para.clear()
        blue_run = desc_para.add_run("When applicable, SME must edit region's Citation Style Guide")
        blue_run.font.color.rgb = RGBColor(0x1D, 0x7A, 0xFC)
        blue_run.bold = True
        blue_run.italic = True
        desc_para.add_run(" Link: ")
        _add_hyperlink(desc_para, "Obligation Drafting / Updates", "file:///C:/confluence/pages/viewpage.action?pageId=2365329841")

        guide_table = doc.add_table(rows=2, cols=2)
        guide_table.rows[0].cells[0].text = "Product Owner"
        owner_para = guide_table.rows[0].cells[1].paragraphs[0]
        owner_para.clear()
        _add_hyperlink(owner_para, "Raut, Divya", "file:///C:/confluence/display/~W620263")
        guide_table.rows[1].cells[0].text = "SME"
        guide_table.rows[1].cells[1].text = "Yiu, Carrie"

        doc.add_heading("Document Structure", level=1)
        levels = doc.add_table(rows=2, cols=8)
        headers = ["Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"]
        for idx, header in enumerate(headers):
            levels.rows[0].cells[idx].text = header
        levels.rows[1].cells[0].text = "2"
        levels.rows[1].cells[1].text = "Title"

        extracted = extract_toc(doc)
        guide = extracted.get("citationStyleGuide") or {}
        self.assertIn("When applicable", guide.get("description", ""))
        self.assertRegex(guide.get("description", ""), r"<(strong|b)>.*When applicable.*</(strong|b)>")
        self.assertRegex(guide.get("description", ""), r"<(em|i)>.*When applicable.*</(em|i)>")
        self.assertEqual(guide["rows"][0]["label"], "Product Owner")
        self.assertIn("Raut, Divya", guide["rows"][0]["value"])
        self.assertIn("file:///C:/confluence/display/~W620263", guide["rows"][0]["value"])
        self.assertEqual(guide["rows"][1]["label"], "SME")

    def test_extract_toc_ignores_scope_and_metadata_rows_inside_citation_guide_block(self):
        doc = Document()
        doc.add_heading("Citation Style Guide Link", level=2)

        guide_table = doc.add_table(rows=1, cols=2)
        guide_table.rows[0].cells[0].text = "Product Owner"
        guide_table.rows[0].cells[1].text = "Raut, Divya"

        polluted_scope_table = doc.add_table(rows=2, cols=2)
        polluted_scope_table.rows[0].cells[0].text = "Document Title"
        polluted_scope_table.rows[0].cells[1].text = "Reference URL"
        polluted_scope_table.rows[1].cells[0].text = "Swiss Ordinance"
        polluted_scope_table.rows[1].cells[1].text = "https://example.com/ordinance"

        polluted_metadata_table = doc.add_table(rows=2, cols=2)
        polluted_metadata_table.rows[0].cells[0].text = "Source Name"
        polluted_metadata_table.rows[0].cells[1].text = "Document Location"
        polluted_metadata_table.rows[1].cells[0].text = "Swiss Federal Council"
        polluted_metadata_table.rows[1].cells[1].text = "Switzerland"

        doc.add_heading("Document Structure", level=1)
        levels = doc.add_table(rows=2, cols=8)
        headers = ["Level", "Name", "Required", "Definition", "Example", "Note", "TOC Requirements", "SME Comments"]
        for idx, header in enumerate(headers):
            levels.rows[0].cells[idx].text = header
        levels.rows[1].cells[0].text = "2"
        levels.rows[1].cells[1].text = "Title"

        extracted = extract_toc(doc)
        guide = extracted.get("citationStyleGuide") or {}
        labels = [row["label"] for row in guide.get("rows") or []]

        self.assertIn("Product Owner", labels)
        self.assertNotIn("Document Title", labels)
        self.assertNotIn("Source Name", labels)

    def test_mhtml_scope_urls_do_not_absorb_evergreen_or_ingestion_values(self):
        mhtml = textwrap.dedent(
            '''\
            MIME-Version: 1.0
            Content-Type: multipart/related; boundary="BOUNDARY"

            --BOUNDARY
            Content-Type: text/html; charset=UTF-8
            Content-Transfer-Encoding: quoted-printable
            Content-Location: file:///C:/exported.html

            <html>
              <body>
                <h2>Scope</h2>
                <table>
                  <tr>
                    <th>Document Title</th>
                    <th>Reference URL</th>
                    <th>Content URL</th>
                    <th>Issuing Authority</th>
                    <th>ASRB ID</th>
                    <th>SME Comments</th>
                    <th>Initial / Evergreen</th>
                    <th>Date of Ingestion</th>
                  </tr>
                  <tr>
                    <td><p><span style="color:#ae2e24;"><strong>Ordonnance test</strong></span></p></td>
                    <td><p><span style="color:#ae2e24;"><a href="https://www.fedlex.admin.ch/eli/cc/2000/243/fr">https://www.fedlex.admin.ch/eli/cc/2000/243/fr</a></span></p></td>
                    <td><p><span style="color:#ae2e24;"><a href="https://www.fedlex.admin.ch/eli/cc/2000/243/fr">https://www.fedlex.admin.ch/eli/cc/2000/243/fr</a></span></p></td>
                    <td><p>Swiss Authority</p></td>
                    <td><p>ASRB-321</p></td>
                    <td><p>This has been repealed and therefore does not need to be ingested.</p></td>
                    <td><p><span style="color:#ae2e24;"><strong>Evergreen</strong></span></p></td>
                    <td><p><span style="color:#ae2e24;"><strong>WIP</strong></span></p></td>
                  </tr>
                </table>
              </body>
            </html>

            --BOUNDARY--
            '''
        )

        with tempfile.NamedTemporaryFile("w", suffix=".doc", delete=False, encoding="utf-8") as handle:
            handle.write(mhtml)
            doc_path = handle.name

        try:
            scope = extract_scope_from_file(doc_path)
        finally:
            Path(doc_path).unlink(missing_ok=True)

        self.assertEqual(len(scope["in_scope"]), 1)
        row = scope["in_scope"][0]
        self.assertIn("https://www.fedlex.admin.ch/eli/cc/2000/243/fr", row["regulator_url"])
        self.assertIn("https://www.fedlex.admin.ch/eli/cc/2000/243/fr", row["content_url"])
        self.assertNotIn("Evergreen", row["regulator_url"])
        self.assertNotIn("WIP", row["content_url"])
        self.assertIn("Evergreen", row["initial_evergreen"])
        self.assertIn("WIP", row["date_of_ingestion"])

if __name__ == "__main__":
    unittest.main()
