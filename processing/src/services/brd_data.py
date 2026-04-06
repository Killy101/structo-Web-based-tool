"""
src/services/brd_data.py

Canonical intermediate representation for a parsed BRD document.

Replaces the ad-hoc dict-passing between extractor.py, metajson_assembler.py,
and pattern_generator.py.  All cleaning and normalization happens once here —
downstream consumers read typed fields, never raw document text.

Usage
-----
    from src.services.brd_data import extract_brd

    brd = extract_brd(docx_path)          # full structured extraction
    brd = extract_brd(docx_path, brd_id)  # + image extraction

    # Pass to pattern generator
    patterns = generate_patterns_from_brd(brd)

    # Pass to metajson assembler
    metajson, filename = assemble_metajson_from_brd(brd, patterns)
"""

from __future__ import annotations

import re
import json
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# 1. Data classes — the canonical contract
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class ScopeEntry:
    """A single in-scope or out-of-scope document."""
    document_title: str
    regulator_url: str = ""
    content_url: str = ""
    content_note: str = ""
    issuing_authority: str = ""
    issuing_authority_code: str = ""
    geography: str = ""
    asrb_id: str = ""
    sme_comments: str = ""
    stable_key: str = ""
    strikethrough: bool = False


@dataclass
class LevelData:
    """
    Merged TOC + citation row for one level.

    All text fields are already cleaned (no \\xa0, no surrounding quotes,
    normalized internal whitespace).  `explicit_patterns` is populated when
    the citation_rules text contains recognizable regex strings — downstream
    pattern generators must treat this as a first-priority override and skip
    their own inference for this level.
    """
    level: int
    name: str = ""
    definition: str = ""
    examples: list[str] = field(default_factory=list)
    required: bool = False
    note: str = ""
    toc_requirements: str = ""
    toc_sme_comments: str = ""
    citation_rules: str = ""        # raw citation rule text, cleaned
    source_of_law: str = ""
    is_citable: str = ""            # "Y" | "N" | ""
    citation_sme_comments: str = ""
    explicit_patterns: list[str] = field(default_factory=list)  # extracted regex, if any
    redjay_xml_tag: str = ""        # from content profile, if available


@dataclass
class ContentProfileLevel:
    """One row from the content profile levels table."""
    level_number: str       # e.g. "Level 3"
    redjay_xml_tag: str
    path: str = ""


@dataclass
class BRDConfigOverrides:
    """
    Explicit config embedded in the BRD document itself (JSON blocks).
    When present these always win over inference.
    """
    level_patterns: dict[str, list[str]] | None = None
    path_transform: dict | None = None
    root_path: str | None = None
    whitespace_handling: dict | None = None
    custom_toc: dict | None = None


@dataclass
class BRDData:
    """
    Complete, normalized representation of one BRD document.

    This is the single object passed between the extraction layer and every
    downstream consumer (pattern generator, metajson assembler, API handlers).
    No consumer should re-parse raw document text.
    """
    # ── Identity ───────────────────────────────────────────────────────────────
    source_file: str = ""
    extracted_at: str = ""
    format: str = "new"                         # "new" | "old"
    language: str = "English"                   # normalized display name
    language_key: str = "english"               # normalized lowercase key

    # ── Metadata ───────────────────────────────────────────────────────────────
    metadata: dict[str, str] = field(default_factory=dict)
    # Canonical keys present regardless of BRD format:
    #   "Content Category Name" | "Source Name"
    #   "Publication Date", "Last Updated Date", "Effective Date",
    #   "Processing Date", "Issuing Agency", "Content URI",
    #   "Geography", "Language", "Delivery Type", "Unique File Id"
    # Old-format extras: "Source Type", "Payload Subtype", "BRD_Version", "Status"

    # ── Scope ──────────────────────────────────────────────────────────────────
    scope_entries: list[ScopeEntry] = field(default_factory=list)   # active (not struck)
    out_of_scope: list[ScopeEntry] = field(default_factory=list)    # struck-through

    # ── Document structure ─────────────────────────────────────────────────────
    levels: list[LevelData] = field(default_factory=list)
    # Sorted ascending by level.level.  Includes only levels >= 2.

    # ── Content profile ────────────────────────────────────────────────────────
    content_profile_levels: list[ContentProfileLevel] = field(default_factory=list)
    heading_annotation: str = "Level 2"
    rc_filename: str = ""           # e.g. "AlabamaAdministrativeCode", "CFR"
    hardcoded_path: str = ""        # e.g. "/us/aladmincode", "/us/cfr"

    # ── Explicit BRD config overrides ─────────────────────────────────────────
    config: BRDConfigOverrides = field(default_factory=BRDConfigOverrides)

    # ── Images (populated only when brd_id is supplied) ───────────────────────
    cell_images: list[dict] = field(default_factory=list)

    # ── Convenience helpers ────────────────────────────────────────────────────

    @property
    def scope_titles(self) -> list[str]:
        """De-duplicated active scope document titles, ready for scope matching."""
        seen: set[str] = set()
        out: list[str] = []
        for e in self.scope_entries:
            t = e.document_title.strip()
            if t and t not in seen:
                seen.add(t)
                out.append(t)
        return out

    @property
    def level_range(self) -> tuple[int, int]:
        """(min_level, max_level) across all levels, defaulting to (2, 7)."""
        nums = [lv.level for lv in self.levels]
        if not nums:
            return (2, 7)
        return (min(nums), max(nums))

    def level_by_num(self, n: int) -> LevelData | None:
        for lv in self.levels:
            if lv.level == n:
                return lv
        return None


# ─────────────────────────────────────────────────────────────────────────────
# 2. Internal helpers — all normalization lives here
# ─────────────────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Strip \\xa0, normalize internal whitespace, remove surrounding quotes."""
    text = text.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = " ".join(text.split())
    return text.strip("\"'\u201c\u201d\u2018\u2019\u00ab\u00bb").strip()


def _clean_multiline(text: str) -> str:
    """Like _clean but preserves intentional newlines (for definition fields)."""
    text = text.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    return "\n".join(line for line in lines if line)


def _normalize_language_key(language: str) -> str:
    key = (language or "").strip().lower().replace("_", "-")
    if any(t in key for t in ["spanish", "español", "espanol", "castellano", "es-"]):
        return "spanish"
    if any(t in key for t in ["portuguese", "português", "portugues", "pt-"]):
        return "portuguese"
    if any(t in key for t in ["chinese", "中文", "汉语", "漢語", "mandarin", "cantonese", "zh"]):
        return "chinese"
    if any(t in key for t in ["japanese", "日本語", "ja-", "jpn"]):
        return "japanese"
    if any(t in key for t in ["korean", "한국어", "ko-", "kor"]):
        return "korean"
    return "english"


def _normalize_level(raw: str) -> int | None:
    """'Level 2' → 2,  '2' → 2,  garbage → None."""
    m = re.search(r"\d+", (raw or "").strip())
    return int(m.group(0)) if m else None


def _required_value(raw: str) -> bool:
    val = (raw or "").strip().lower()
    return val in ("true", "yes", "y", "1")


def _split_examples(raw: str) -> list[str]:
    """
    Split a BRD example cell into individual example strings.

    Many BRDs (especially CFR) pack multiple example types into one cell,
    separated by newlines or pipe characters.  Each entry may itself contain
    both a structural identifier and a definition-of-terms label:

        "(b) | \"Financial end user\""
        "(2) | \"Seller's interest\""

    We split on newlines first (each newline = a separate example entry),
    then on " | " within each entry.  The full split items are all preserved
    in the returned list so that pattern inference can inspect them.

    Callers that only want the PRIMARY (structural) token per entry should
    use _split_primary_examples() on the result.
    """
    # Normalize non-breaking spaces and carriage returns
    s = raw.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n").strip()

    # Strip trailing boilerplate
    for suffix in ("; etc.", ", etc.", " etc."):
        if s.lower().endswith(suffix):
            s = s[: -len(suffix)].strip()

    if not s:
        return []

    # Split on newlines → individual entries
    entries: list[str] = []
    for line in s.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Within each line, split on " | " to get sub-examples
        parts = [p.strip().strip("\"'\u201c\u201d") for p in line.split(" | ") if p.strip()]
        entries.extend(p for p in parts if p)

    # Fallback: semicolon-separated on a single line
    if not entries and ";" in s:
        entries = [t.strip().strip("\"'\u201c\u201d") for t in s.split(";") if t.strip()]

    return entries if entries else ([s.strip("\"'\u201c\u201d")] if s else [])


def _split_primary_examples(examples: list[str]) -> list[str]:
    """
    From a list of raw example strings (which may contain both a structural
    identifier AND a definition-of-terms label separated by ' | ' or newline),
    return only the FIRST token per entry — the structural identifier.

    This prevents definition-of-terms labels like '"Financial end user"' from
    being fed to the keyword heuristic and generating wrong patterns.
    """
    result: list[str] = []
    for ex in examples:
        if not ex or not ex.strip():
            continue
        # Each raw example may be "structural | label" — take only the first part
        first = re.split(r'\s*[\|\n]\s*', ex.strip())[0].strip().strip('"')
        if first:
            result.append(first)
    return result


# ── Regex extraction from citation rule text ──────────────────────────────────

def _looks_regex(raw: str) -> bool:
    lowered = raw.lower()
    if "<level" in lowered or "example:" in lowered:
        return False
    if re.search(r"\+\s*\"", raw):
        return False
    return bool(
        re.search(r"(\^|\$|\\[dDsSwWbBAZz]|\\\\.)", raw)
        or re.search(r"\[[^\]]+\](?:\{\d+(?:,\d*)?\}|[+*?])?", raw)
        or re.search(r"\([^)]*\|[^)]*\)", raw)
        or re.search(r"(?:\)|\]|\.|[A-Za-z0-9])[+*?]", raw)
    )


def _infer_heading_regex(candidate: str) -> str | None:
    """Convert a plain heading example ('Chapter 1') to a regex string."""
    candidate = re.sub(r"\s+", " ", candidate).strip(" \"'`")
    candidate = re.sub(r"^level\s*\d+\s*", "", candidate, flags=re.IGNORECASE).strip(" :-")
    candidate = re.sub(
        r"^(example|examples|pattern|regex|rule)\s*:\s*", "", candidate, flags=re.IGNORECASE
    ).strip()
    if not candidate:
        return None

    if "§" in candidate and re.search(r"[0-9A-Za-z]", candidate):
        num = r"[0-9A-Za-z]+(?:[-.][0-9A-Za-z]+)*"
        return rf"^(?:SECTION|Section|Sec\.?\s*)?\s*§{{1,2}}\s*{num}(?:\([0-9A-Za-z]+\))*$"

    m = re.match(
        r"^(chapter|part|division|subdivision|section|article|rule|title|subtitle"
        r"|subpart|subchapter|appendix|schedule|exhibit|attachment|form)"
        r"\s+([0-9]+(?:[A-Z])?(?:[-.][0-9A-Z]+)*)$",
        candidate,
        flags=re.IGNORECASE,
    )
    if not m:
        return None

    keyword = m.group(1)
    kw = (
        r"(ARTICLE|Article|Art\.?)"
        if keyword.lower() == "article"
        else f"({keyword.upper()}|{keyword.title()})"
    )
    return f"^{kw} ?[0-9]+[A-Z]?(?:[-.][0-9A-Z]+)*$"


def _extract_explicit_patterns(text: str) -> list[str]:
    """
    Extract explicit regex strings from free-text citation rule blocks.
    Returns an empty list when no valid regex is found — the pattern generator
    should then do its own inference from definition + examples.
    """
    if not text or not text.strip():
        return []

    lines = text.replace("\r", "\n").replace("\t", " ").split("\n")
    cleaned: list[str] = []
    plain_candidates: list[str] = []

    for line in lines:
        cur = line.strip()
        if not cur:
            continue
        cur = re.sub(r"^<\s*level\s*\d+\s*>\s*", "", cur, flags=re.IGNORECASE).strip()
        cur = re.sub(r"^level\s*\d+\s*[:\-]?\s*", "", cur, flags=re.IGNORECASE).strip()
        cur = re.sub(r"^[-*\d.)\s]+", "", cur).strip()
        cur = re.sub(
            r"^(pattern|regex|rule|example|examples|notes?)\s*:\s*", "", cur,
            flags=re.IGNORECASE,
        ).strip()
        cur = cur.strip('"\'`').strip().rstrip(",")
        cur = re.sub(r"\bexample\s*:.*$", "", cur, flags=re.IGNORECASE).strip()
        cur = re.sub(r"\*\s*note\s*:.*$", "", cur, flags=re.IGNORECASE).strip()
        if not cur:
            continue
        slash = re.match(r"^/(.+)/[gimsuy]*$", cur)
        if slash:
            cur = slash.group(1)
        if _looks_regex(cur):
            cleaned.append(cur)
            continue
        for segment in re.split(r"[;\n]+", cur):
            s = segment.strip()
            if s:
                plain_candidates.append(s)

    for candidate in plain_candidates:
        inferred = _infer_heading_regex(candidate)
        if inferred:
            cleaned.append(inferred)

    # Deduplicate preserving order
    seen: set[str] = set()
    deduped: list[str] = []
    for p in cleaned:
        if p not in seen:
            seen.add(p)
            deduped.append(p)
    return deduped


# ── BRD config block extraction (JSON embedded in document text) ──────────────

def _collect_document_text(doc) -> str:
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    lines.append(t)
    return "\n".join(lines)


def _cleanup_json_like(raw: str) -> str:
    raw = (
        raw.replace("\u201c", '"').replace("\u201d", '"')
           .replace("\u2018", "'").replace("\u2019", "'")
           .replace("\uff1a", ":")
    )
    cleaned = re.sub(r"//.*?$", "", raw, flags=re.MULTILINE)
    cleaned = re.sub(r"/\*.*?\*/", "", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r",\s*([}\]])", r"\1", cleaned)
    return cleaned.strip()


def _extract_braced_block(text: str, start: int) -> str | None:
    if start < 0 or start >= len(text) or text[start] != "{":
        return None
    depth = 0
    in_string = False
    quote_char = ""
    escaped = False
    for i in range(start, len(text)):
        ch = text[i]
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == quote_char:
                in_string = False
            continue
        if ch in ('"', "'"):
            in_string = True
            quote_char = ch
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return None


def _extract_bracket_block(text: str, start: int) -> str | None:
    if start < 0 or start >= len(text) or text[start] != "[":
        return None
    depth = 0
    in_string = False
    quote_char = ""
    escaped = False
    for i in range(start, len(text)):
        ch = text[i]
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == quote_char:
                in_string = False
            continue
        if ch in ('"', "'"):
            in_string = True
            quote_char = ch
            continue
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return None


def _parse_json_block(raw: str) -> dict | None:
    cleaned = _cleanup_json_like(raw)
    try:
        parsed = json.loads(cleaned)
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        return None


def _extract_config_object(text: str, *keys: str) -> dict | None:
    for key in keys:
        m = re.search(
            rf"(?:\"{re.escape(key)}\"|'{re.escape(key)}'|\b{re.escape(key)}\b)\s*:\s*{{",
            text,
        )
        if not m:
            continue
        brace_idx = text.find("{", m.end() - 1)
        if brace_idx < 0:
            continue
        block = _extract_braced_block(text, brace_idx)
        if not block:
            continue
        parsed = _parse_json_block(block)
        if isinstance(parsed, dict) and parsed:
            return parsed
    return None


def _extract_config_scalar(text: str, *keys: str) -> str | None:
    for key in keys:
        m = re.search(
            rf"(?:\"{re.escape(key)}\"|'{re.escape(key)}'|\b{re.escape(key)}\b)\s*:\s*\"([^\"]+)\"",
            text,
        )
        if m and m.group(1).strip():
            return m.group(1).strip()
    return None


def _extract_path_transform_relaxed(text: str) -> dict | None:
    """Parse pathTransform even when BRD text is not strict JSON."""
    cleaned = _cleanup_json_like(text)
    marker = re.search(
        r"(?is)(?:\"pathTransform\"|'pathTransform'|pathTransform"
        r"|path_transform|\"path_transform\"|'path_transform')\s*:\s*{",
        cleaned,
    )
    if not marker:
        return None

    root_open = cleaned.find("{", marker.end() - 1)
    if root_open < 0:
        return None
    root_block = _extract_braced_block(cleaned, root_open)
    if not root_block:
        return None

    strict = _parse_json_block(root_block)
    if isinstance(strict, dict) and strict:
        return strict

    # Manual level-by-level parse when strict JSON fails
    out: dict[str, dict] = {}
    for level_match in re.finditer(r"(?is)['\"]?(\d{1,2})['\"]?\s*:\s*{", root_block):
        level = level_match.group(1)
        level_open = root_block.find("{", level_match.end() - 1)
        if level_open < 0:
            continue
        level_block = _extract_braced_block(root_block, level_open)
        if not level_block:
            continue

        pat_match = re.search(
            r"(?is)(?:\"patterns\"|'patterns'|patterns)\s*:\s*\[", level_block
        )
        if not pat_match:
            continue
        pat_open = level_block.find("[", pat_match.end() - 1)
        if pat_open < 0:
            continue
        pat_block = _extract_bracket_block(level_block, pat_open)
        if not pat_block:
            continue

        patterns: list[list] = []
        # Extract each [find, replace, flag, extra] row
        for item_m in re.finditer(
            r"\[\s*(['\"])(.*?)\1\s*,\s*(['\"])(.*?)\3\s*,\s*(-?\d+)\s*,\s*(['\"])(.*?)\6\s*\]",
            pat_block,
            flags=re.DOTALL,
        ):
            find    = item_m.group(2).replace(r"\"", '"').replace(r"\\'", "'").strip()
            replace = item_m.group(4).replace(r"\"", '"').replace(r"\\'", "'").strip()
            flag    = int(item_m.group(5))
            extra   = item_m.group(7).strip()
            patterns.append([find, replace, flag, extra])

        case_m = re.search(
            r"(?is)(?:\"case\"|'case'|case)\s*:\s*(['\"])(.*?)\1", level_block
        )
        case_val = case_m.group(2) if case_m else ""

        if patterns:
            out[level] = {"patterns": patterns, "case": case_val}

    return out or None


def _extract_brd_config_overrides(doc_text: str) -> BRDConfigOverrides:
    """
    Parse any JSON config blocks embedded in the BRD document text.
    These are the highest-priority source for patterns and path transforms.
    """
    overrides = BRDConfigOverrides()

    path_transform = (
        _extract_path_transform_relaxed(doc_text)
        or _extract_config_object(doc_text, "pathTransform", "path_transform")
    )
    if isinstance(path_transform, dict) and path_transform:
        overrides.path_transform = path_transform

    level_patterns = _extract_config_object(doc_text, "levelPatterns", "level_patterns")
    if isinstance(level_patterns, dict) and level_patterns:
        overrides.level_patterns = {str(k): v for k, v in level_patterns.items()}

    whitespace = _extract_config_object(
        doc_text, "whitespaceHandling", "whitespace_handling"
    )
    if isinstance(whitespace, dict) and whitespace:
        overrides.whitespace_handling = whitespace

    custom_toc = _extract_config_object(doc_text, "custom_toc", "customToc")
    if isinstance(custom_toc, dict) and custom_toc:
        overrides.custom_toc = custom_toc

    root_path = _extract_config_scalar(doc_text, "rootPath", "root_path")
    if root_path:
        overrides.root_path = root_path

    return overrides


# ── Section-level extraction delegates ───────────────────────────────────────

def _extract_scope_entries(doc) -> tuple[list[ScopeEntry], list[ScopeEntry]]:
    """
    Delegates to the existing scope extractor, then converts its dict output
    to ScopeEntry dataclasses with cleaned fields.
    """
    from .extractors.scope_extractor import extract_scope

    raw = extract_scope(doc)
    active: list[ScopeEntry] = []
    struck: list[ScopeEntry] = []

    def _to_entry(d: dict) -> ScopeEntry:
        return ScopeEntry(
            document_title=_clean(d.get("document_title", "")),
            regulator_url=_clean(d.get("regulator_url", "")),
            content_url=_clean(d.get("content_url", "")),
            content_note=_clean(d.get("content_note", "")),
            issuing_authority=_clean(d.get("issuing_authority", "")),
            issuing_authority_code=_clean(d.get("issuing_authority_code", "")),
            geography=_clean(d.get("geography", "")),
            asrb_id=_clean(d.get("asrb_id", "")),
            sme_comments=_clean(d.get("sme_comments", "")),
            stable_key=_clean(d.get("stable_key", d.get("stableKey", ""))),
            strikethrough=bool(d.get("strikethrough", False)),
        )

    for entry in raw.get("in_scope", []):
        if isinstance(entry, dict):
            active.append(_to_entry(entry))

    for entry in raw.get("out_of_scope", []):
        if isinstance(entry, dict):
            struck.append(_to_entry(entry))

    return active, struck


def _extract_metadata_normalized(doc, format_: str) -> dict[str, str]:
    """
    Delegates to the existing metadata extractor, then maps its raw output to
    canonical metajson keys in one place.  Output always uses the same key set
    regardless of BRD format — no more dual key-alias lookups downstream.
    """
    from .extractors.metadata_extractor import extract_metadata

    raw = extract_metadata(doc)

    def t(key: str) -> str:
        v = raw.get(key)
        return _clean(str(v)) if v else ""

    detected_format = raw.get("_format", format_)

    if detected_format == "new":
        return {
            "_format": "new",
            "Content Category Name":     t("content_category_name") or t("document_title"),
            "Authoritative Source":      t("authoritative_source") or t("issuing_agency"),
            "Content Type":              t("content_type"),
            "Publication Date":          t("publication_date")  or "{iso-date}",
            "Last Updated Date":         t("last_updated_date") or "{iso-date}",
            "Effective Date":            t("effective_date") or "{iso-date}",
            "Comment Due Date":          t("comment_due_date"),
            "Compliance Date":           t("compliance_date"),
            "Processing Date":           t("processing_date") or "{iso-date}",
            "Name":                      t("name") or t("document_title"),
            "Issuing Agency":            t("issuing_agency"),
            "Related Government Agency": t("related_government_agency"),
            "Content URI":               t("content_uri") or "{string}",
            "Geography":                 t("geography"),
            "Language":                  t("language"),
            "Impacted Citation":         t("impacted_citation"),
            "Payload Type":              t("payload_type"),
            "Payload Subtype":           t("payload_subtype"),
            "Summary":                   t("summary"),
            "SME Comments":              t("sme_comments"),
            "Status":                    t("status"),
            "Delivery Type":             t("delivery_type") or "{string}",
            "Unique File Id":            "{string}",
        }
    else:
        return {
            "_format": "old",
            "Authoritative Source":      t("authoritative_source") or t("issuing_agency"),
            "Source Name":               t("content_category_name") or t("document_title"),
            "Source Type":               t("source_type"),
            "Content Type":              t("content_type"),
            "Publication Date":          t("publication_date") or "{iso-date}",
            "Last Updated Date":         t("last_updated_date") or "{iso-date}",
            "Effective Date":            t("effective_date") or "{iso-date}",
            "Comment Due Date":          t("comment_due_date"),
            "Compliance Date":           t("compliance_date"),
            "Processing Date":           t("processing_date") or "{iso-date}",
            "Name":                      t("name") or t("document_title"),
            "Issuing Agency":            t("issuing_agency"),
            "Related Government Agency": t("related_government_agency"),
            "Content URI":               t("content_uri") or "{string}",
            "Geography":                 t("geography"),
            "Language":                  t("language"),
            "Impacted Citation":         t("impacted_citation"),
            "Payload Type":              t("payload_type"),
            "Payload Subtype":           t("payload_subtype") or "Acts",
            "Summary":                   t("summary"),
            "SME Comments":              t("sme_comments"),
            "Status":                    t("status") or "Effective",
            "BRD_Version":               t("version"),
            "Delivery Type":             t("delivery_type") or "{string}",
            "Unique File Id":            "{string}",
        }


def _examples_from_redjay_tag(tag: str) -> list[str]:
    """
    Extract <title> values from a redjayXmlTag string.

    Example input:
        '<section level="8"><title>§ 217.132</title></section>'
    Returns: ['§ 217.132']

    These are real identifier samples from the document and are the most
    reliable source for pattern inference — more reliable than the BRD
    example cell which often mixes structural ids with definition labels.
    """
    if not tag or "hardcoded" in tag.lower():
        return []
    return [
        m.group(1).strip()
        for m in re.finditer(r"<title>(.*?)</title>", tag, re.IGNORECASE | re.DOTALL)
        if m.group(1).strip()
    ]


def _examples_from_citation_rule(citation_rules: str) -> list[str]:
    """
    Extract concrete example identifiers from a citation rule string.

    Citation rules often contain an 'Example:' clause with real identifiers:
        '<level 2> *number only* + " C.F.R. " + <Level 8>
         Example: "12 C.F.R. § 217.132"'

    We extract the part after "Example:" and then pull out the per-level
    tokens.  For most levels, the last quoted token IS the level identifier.

    Returns a list of candidate example strings — the caller should then
    pick the most useful ones for pattern inference.
    """
    if not citation_rules:
        return []

    examples: list[str] = []

    # Find all "Example: ..." clauses
    for m in re.finditer(r'[Ee]xample\s*:\s*"?([^"\n]+)"?', citation_rules):
        raw = m.group(1).strip().strip('"')
        if raw:
            examples.append(raw)

    return examples


def _build_levels(doc) -> list[LevelData]:
    """
    Merge TOC and citation tables into a single sorted list of LevelData.

    Source priority for pattern inference (highest to lowest):
      1. redjayXmlTag <title> values  — real document identifiers, most accurate
      2. TOC example cell (primary token only)  — structural id before any '|'
      3. Citation rule Example: clause  — identifier embedded in citation format
      4. TOC definition text  — keyword description ("\"Chapter\" + roman numeral")

    All sources are stored on LevelData so the pattern generator can see
    the full context.  explicit_patterns (regex extracted from citation text)
    are populated here once — no downstream code re-parses citation text.
    """
    from .extractors.toc_extractor import extract_toc
    from .extractors.citations_extractor import extract_citations

    toc_raw = extract_toc(doc)
    citations_raw = extract_citations(doc)

    # Build citation index: level_str → citation row dict
    citation_index: dict[str, dict] = {}
    for ref in citations_raw.get("references", []):
        if not isinstance(ref, dict):
            continue
        lvl_str = str(ref.get("level", "")).strip()
        if lvl_str:
            citation_index[lvl_str] = ref

    levels: list[LevelData] = []
    seen: set[int] = set()

    for section in toc_raw.get("sections", []):
        if not isinstance(section, dict):
            continue

        raw_level = str(section.get("level") or section.get("id") or "")
        level_num = _normalize_level(raw_level)
        if level_num is None or level_num < 2:
            continue
        if level_num in seen:
            continue
        seen.add(level_num)

        name = _clean(str(section.get("name", "")))
        definition = _clean_multiline(str(section.get("definition", "")))
        raw_example = str(section.get("example", ""))
        required = _required_value(str(section.get("required", "")))
        note = _clean_multiline(str(section.get("note") or ""))
        toc_requirements = _clean_multiline(str(
            section.get("tocRequirements") or section.get("toc_requirements") or ""
        ))
        toc_sme_comments = _clean_multiline(str(
            section.get("smeComments") or section.get("sme_comments") or ""
        ))

        # Pull citation row
        cit = citation_index.get(str(level_num), {})
        citation_rules = _clean_multiline(str(
            cit.get("citationRules") or cit.get("citation_rules") or ""
        ))
        source_of_law = _clean(str(cit.get("sourceOfLaw") or cit.get("source_of_law") or ""))
        is_citable = _clean(str(cit.get("isCitable") or cit.get("is_citable") or ""))
        citation_sme_comments = _clean_multiline(str(
            cit.get("smeComments") or cit.get("sme_comments") or ""
        ))

        # ── Build examples with priority ordering ─────────────────────────────
        # Start with raw TOC examples (preserving newlines for multi-type levels)
        toc_examples = _split_examples(raw_example)

        # Supplement with citation rule examples when they contain real identifiers
        # (citation rule examples are more precise for § and parenthetical levels)
        cit_examples = _examples_from_citation_rule(citation_rules)

        # Use TOC examples as primary; citation examples fill gaps
        # The generator's _split_primary_examples() will take only the
        # first structural token from each entry.
        combined_examples = toc_examples or cit_examples

        # Extract explicit regex from citation rules — once, right here
        explicit_patterns = _extract_explicit_patterns(citation_rules)

        levels.append(
            LevelData(
                level=level_num,
                name=name,
                definition=definition,
                examples=combined_examples,
                required=required,
                note=note,
                toc_requirements=toc_requirements,
                toc_sme_comments=toc_sme_comments,
                citation_rules=citation_rules,
                source_of_law=source_of_law,
                is_citable=is_citable,
                citation_sme_comments=citation_sme_comments,
                explicit_patterns=explicit_patterns,
            )
        )

    # Also include levels that appear in citations but not in TOC
    for lvl_str, cit in citation_index.items():
        level_num = _normalize_level(lvl_str)
        if level_num is None or level_num < 2 or level_num in seen:
            continue
        seen.add(level_num)

        citation_rules = _clean_multiline(str(
            cit.get("citationRules") or cit.get("citation_rules") or ""
        ))
        cit_examples   = _examples_from_citation_rule(citation_rules)
        explicit_patterns = _extract_explicit_patterns(citation_rules)

        levels.append(
            LevelData(
                level=level_num,
                examples=cit_examples,
                citation_rules=citation_rules,
                source_of_law=_clean(str(cit.get("sourceOfLaw") or cit.get("source_of_law") or "")),
                is_citable=_clean(str(cit.get("isCitable") or cit.get("is_citable") or "")),
                citation_sme_comments=_clean_multiline(str(cit.get("smeComments") or cit.get("sme_comments") or "")),
                explicit_patterns=explicit_patterns,
            )
        )

    levels.sort(key=lambda lv: lv.level)
    return levels


def _build_content_profile_levels(
    doc, levels: list[LevelData]
) -> tuple[list[ContentProfileLevel], str, str]:
    """
    Build ContentProfileLevel rows from the content profile extractor.

    Returns (cp_levels, rc_filename, hardcoded_path).

    Side-effects on LevelData (both applied before pattern generation):
      1. Sets LevelData.redjay_xml_tag from the content profile row.
      2. PREPENDS redjayXmlTag <title> values to LevelData.examples.
         These are real identifiers from actual documents and are the most
         reliable signal for pattern inference — more accurate than the raw
         BRD example cell which mixes structural ids with prose labels.
    """
    from .extractors.content_profile_extractor import extract_content_profile

    try:
        cp = extract_content_profile(doc)
    except Exception:
        cp = {}

    cp_level_rows = cp.get("levels", [])
    cp_levels: list[ContentProfileLevel] = []

    level_map = {lv.level: lv for lv in levels}

    for row in cp_level_rows:
        if not isinstance(row, dict):
            continue
        level_num_raw = str(row.get("levelNumber") or row.get("level_number") or "")
        m = re.search(r"\d+", level_num_raw)
        if not m:
            continue
        level_num = int(m.group(0))
        tag  = str(row.get("redjayXmlTag") or row.get("redjay_xml_tag") or "").strip()
        path = str(row.get("path") or "").strip()

        cp_levels.append(ContentProfileLevel(
            level_number=f"Level {level_num}",
            redjay_xml_tag=tag,
            path=path,
        ))

        if level_num in level_map:
            lv = level_map[level_num]

            # 1. Store the tag string
            if tag:
                lv.redjay_xml_tag = tag

            # 2. Prepend redjay <title> examples — highest priority signal
            redjay_examples = _examples_from_redjay_tag(tag)
            if redjay_examples:
                existing = lv.examples or []
                seen_ex: set[str] = set(redjay_examples)
                extras = [e for e in existing if e not in seen_ex]
                lv.examples = redjay_examples + extras

    rc_filename    = str(cp.get("rc_filename")    or "").strip()
    hardcoded_path = str(cp.get("hardcoded_path") or "").strip()

    return cp_levels, rc_filename, hardcoded_path


# ─────────────────────────────────────────────────────────────────────────────
# 3. Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_brd(docx_path: str, brd_id: str | None = None) -> BRDData:
    """
    Single-pass extraction from a .docx BRD file.

    All cleaning, normalization, and merging happens here.
    The returned BRDData object is the only thing downstream code needs.

    Parameters
    ----------
    docx_path : str
        Absolute or relative path to the .docx file.
    brd_id : str | None
        When provided, also extracts cell images.  The returned BRDData
        will have cell_images populated.

    Returns
    -------
    BRDData
        Fully normalized, merged representation of the BRD.
    """
    doc = Document(docx_path)
    doc_text = _collect_document_text(doc)

    # ── 1. Metadata ────────────────────────────────────────────────────────────
    metadata = _extract_metadata_normalized(doc, "new")
    format_  = metadata.get("_format", "new")

    # ── 2. Language (resolve once) ─────────────────────────────────────────────
    raw_language = (
        metadata.get("Language")
        or metadata.get("language")
        or "English"
    )
    language_key = _normalize_language_key(raw_language)
    # Produce a canonical display name
    _display = {
        "spanish": "Spanish", "portuguese": "Portuguese", "chinese": "Chinese",
        "japanese": "Japanese", "korean": "Korean", "english": "English",
    }
    language_display = _display.get(language_key, raw_language)

    # ── 3. Scope ───────────────────────────────────────────────────────────────
    scope_entries, out_of_scope = _extract_scope_entries(doc)

    # ── 4. Levels (TOC + citations merged) ────────────────────────────────────
    levels = _build_levels(doc)

    # ── 5. Content profile (enriches LevelData.redjay_xml_tag as side-effect) ─
    cp_levels, cp_rc_filename, cp_hardcoded_path = _build_content_profile_levels(doc, levels)

    # ── 6. Embedded BRD config overrides ──────────────────────────────────────
    config = _extract_brd_config_overrides(doc_text)

    # ── 7. Images ──────────────────────────────────────────────────────────────
    cell_images: list[dict] = []
    if brd_id:
        try:
            from .extractors.image_extractor import extract_and_store_images
            cell_images = extract_and_store_images(doc, docx_path, brd_id=brd_id)
        except Exception as exc:
            print(f"[WARN brd_data] Image extraction failed: {exc}")

    return BRDData(
        source_file=Path(docx_path).name,
        extracted_at=datetime.utcnow().isoformat() + "Z",
        format=format_,
        language=language_display,
        language_key=language_key,
        metadata=metadata,
        scope_entries=scope_entries,
        out_of_scope=out_of_scope,
        levels=levels,
        content_profile_levels=cp_levels,
        heading_annotation="Level 2",
        rc_filename=cp_rc_filename,
        hardcoded_path=cp_hardcoded_path,
        config=config,
        cell_images=cell_images,
    )


# ─────────────────────────────────────────────────────────────────────────────
# 4. Adapter: BRDData → existing generate_level_patterns() input format
# ─────────────────────────────────────────────────────────────────────────────

def brd_to_level_pattern_input(brd: BRDData) -> list[dict[str, Any]]:
    """
    Convert BRDData.levels to the list[dict] format consumed by
    generate_level_patterns() in pattern_generator/__init__.py.

    Levels that already have explicit_patterns are included with their
    explicit patterns pre-loaded so the generator's first-priority check
    (citation text override) fires immediately without re-parsing.
    """
    result: list[dict[str, Any]] = []
    for lv in brd.levels:
        entry: dict[str, Any] = {
            "level":      lv.level,
            "definition": lv.definition or lv.citation_rules,
            "examples":   lv.examples,
            "required":   lv.required,
            "name":       lv.name,
        }
        if lv.redjay_xml_tag:
            entry["redjayXmlTag"] = lv.redjay_xml_tag
        result.append(entry)
    return result


def brd_to_metajson_input(brd: BRDData) -> dict[str, Any]:
    """
    Convert BRDData to the payload shape expected by assemble_metajson()
    in metajson_assembler.py, for backward compatibility during migration.

    Once assemble_metajson() is updated to accept BRDData directly this
    adapter can be deleted.
    """
    scope_dict = {
        "in_scope": [
            {
                "document_title": e.document_title,
                "regulator_url":  e.regulator_url,
                "content_url":    e.content_url,
                "content_note":   e.content_note,
                "issuing_authority": e.issuing_authority,
                "issuing_authority_code": e.issuing_authority_code,
                "geography":      e.geography,
                "asrb_id":        e.asrb_id,
                "sme_comments":   e.sme_comments,
                "stable_key":     e.stable_key,
                "strikethrough":  False,
            }
            for e in brd.scope_entries
        ],
        "out_of_scope": [
            {
                "document_title": e.document_title,
                "strikethrough":  True,
            }
            for e in brd.out_of_scope
        ],
    }

    citations_dict = {
        "references": [
            {
                "level":         str(lv.level),
                "citationRules": lv.citation_rules,
                "sourceOfLaw":   lv.source_of_law,
                "isCitable":     lv.is_citable,
                "smeComments":   lv.citation_sme_comments,
            }
            for lv in brd.levels
        ]
    }

    content_profile_dict: dict[str, Any] = {
        "heading_annotation": brd.heading_annotation,
        "rc_filename":        brd.rc_filename,
        "hardcoded_path":     brd.hardcoded_path,
        "levels": [
            {
                "levelNumber":  cp.level_number,
                "redjayXmlTag": cp.redjay_xml_tag,
                "path":         cp.path,
            }
            for cp in brd.content_profile_levels
        ],
    }

    toc_dict = {
        "sections": [
            {
                "id":              str(lv.level),
                "level":           str(lv.level),
                "name":            lv.name,
                "required":        "Yes" if lv.required else "No",
                "definition":      lv.definition,
                "example":         "; ".join(lv.examples),
                "note":            lv.note,
                "tocRequirements": lv.toc_requirements,
                "smeComments":     lv.toc_sme_comments,
            }
            for lv in brd.levels
        ]
    }

    brd_config: dict[str, Any] = {}
    if brd.config.level_patterns:
        brd_config["levelPatterns"] = brd.config.level_patterns
    if brd.config.path_transform:
        brd_config["pathTransform"] = brd.config.path_transform
    if brd.config.root_path:
        brd_config["rootPath"] = brd.config.root_path
    if brd.config.whitespace_handling:
        brd_config["whitespaceHandling"] = brd.config.whitespace_handling

    return {
        "metadata":        brd.metadata,
        "language":        brd.language,
        "scope":           scope_dict,
        "citations":       citations_dict,
        "toc":             toc_dict,
        "contentProfile":  content_profile_dict,
        "brdConfig":       brd_config or None,
        "format":          brd.format,
        "cell_images":     brd.cell_images,
        "extracted_at":    brd.extracted_at,
        "source_file":     brd.source_file,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 5. Explicit-patterns override applier
#    Call this AFTER generate_level_patterns() to inject explicit overrides
#    for any level where LevelData.explicit_patterns is non-empty.
# ─────────────────────────────────────────────────────────────────────────────

def apply_explicit_pattern_overrides(
    inferred: dict[str, list[str]],
    brd: BRDData,
) -> dict[str, list[str]]:
    """
    For every level that has explicit_patterns extracted from its citation_rules,
    replace the inferred patterns with the explicit ones.

    This is the single authoritative place where citation-text regex overrides
    are applied — previously this logic was duplicated in generate_level_patterns()
    and in process.py.

    Call pattern:
        patterns = generate_level_patterns(language=brd.language, levels=level_input)
        patterns = apply_explicit_pattern_overrides(patterns, brd)
    """
    result = dict(inferred)
    for lv in brd.levels:
        if lv.explicit_patterns:
            result[str(lv.level)] = list(lv.explicit_patterns)
    # Level 2 is always catch-all
    result["2"] = [r"^.*$"]
    return result