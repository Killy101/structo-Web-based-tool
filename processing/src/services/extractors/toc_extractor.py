"""
src/services/extractors/toc_extractor.py

Extracts the Document Structure / TOC levels table from a BRD .docx file.

The TOC table has 8 columns:
  Level | Name | Required | Definition | Example | Note | TOC Requirements | SME Comments

This maps directly to the frontend TocRow interface:
  { id, level, name, required, definition, example, note, tocRequirements, smeComments }
"""

import html as html_lib
import re
from typing import Any

# Helpers

# ─────────────────────────────────────────────
# Shared: parse hardcoded path from definition
# ─────────────────────────────────────────────

# Matches a path segment in a definition, e.g.:
#   "Hardcoded – /us"                → "/us"
#   "Hardcoded – \"/cl\""            → "/cl"
#   "Hardcoded - (/de)"              → "/de"
#   "Hardcoded: /kr"                 → "/kr"
_HARDCODED_PATH_RE = re.compile(
    r"hardcoded(?:\s+path)?\s*[–\-—:=]?\s*[\"'“”‘’(\[]*\s*(/[A-Za-z0-9][A-Za-z0-9_./-]*)",
    re.IGNORECASE,
)

# Fallback: capture the first slash-prefixed token anywhere in the definition.
_ANY_PATH_RE = re.compile(r"(/[A-Za-z0-9][A-Za-z0-9_./-]*)")

# Matches a definition that is *just* a path, e.g. "/kr" or '"/KRNARKActs"'.
_BARE_PATH_RE = re.compile(r"^\s*[\"'“”‘’(\[]*\s*(/[A-Za-z][A-Za-z0-9_./-]*)\s*[\"'“”‘’)\]]*\s*$")


def _extract_path_from_definition(definition: str) -> str:
    """
    Pull the path segment out of a Level 0 / Level 1 definition.
    Handles quoted values and lightly localized wording as long as a
    slash-prefixed path is present somewhere in the text.
    """
    cleaned = (definition or "").strip()
    if not cleaned:
        return ""

    for pattern in (_HARDCODED_PATH_RE, _ANY_PATH_RE):
        match = pattern.search(cleaned)
        if match:
            return match.group(1).rstrip(".,;:)]}\"'”’")

    # Fallback: definition is itself a bare path segment (e.g. KR.NARK Level 0/1)
    bare = _BARE_PATH_RE.match(cleaned)
    return bare.group(1).rstrip(".,;:)]}\"'”’") if bare else ""


def _normalize_level_value(raw: str) -> str:
    """Normalize level cell values like `0`, `L0`, or `Level 0` → `0`."""
    text = _clean(raw)
    if not text:
        return ""

    digit_match = re.search(r"\b(?:level|lvl|l)?\s*(\d+)\b", text, re.IGNORECASE)
    if digit_match:
        return digit_match.group(1)

    short = text.strip()
    compact_match = re.fullmatch(r"[A-Za-z]?(\d{1,2})", short)
    if compact_match:
        return compact_match.group(1)

    return short


def _clean(text: str) -> str:
    """Normalise whitespace while preserving meaningful line breaks."""
    normalized = text.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def _clean_rich_text(text: str) -> str:
    normalized = text.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\s*(<br\s*/?>)\s*", r"\1", normalized, flags=re.IGNORECASE)
    return normalized.strip()


def _normalize_heading(text: str) -> str:
    normalized = _clean(text).lower().replace("\u00a0", " ")
    normalized = normalized.replace("*", " ").replace("{", " ").replace("}", " ")
    normalized = re.sub(r"^[\s\u2022•·\-–—#:]+", "", normalized)
    normalized = re.sub(r"\(\s*#.*?\)", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip(" :-")


_SECTION_HEADING_PATTERNS: tuple[str, ...] = (
    r"^metadata(?:\s+details?)?(?:\s*\([^)]*\))?$",
    r"^details?(?:\s*\([^)]*\))?$",
    r"^exceptions?(?:\s*\([^)]*\))?$",
    r"^updates?(?:\s*\([^)]*\))?$",
    r"^scope(?:\s*\([^)]*\))?$",
    r"^how to identify(?:\s*\([^)]*\))?$",
    r"^frequency of updates(?:\s*\([^)]*\))?$",
    r"^(?:content category|source) to be monitored for updates.*$",
    r"^file delivery requirements(?:\s*\([^)]*\))?$",
    r"^file separation(?:\s*\([^)]*\))?$",
    r"^file naming conventions(?:\s*\([^)]*\))?$",
    r"^rc file naming conventions(?:\s*\([^)]*\))?$",
    r"^zip file naming conventions(?:\s*\([^)]*\))?$",
    r"^brd signed[- ]off by sme(?:\s*\([^)]*\))?$",
    r"^(?:toc\s+with\s+)?document structure(?:\s+levels?)?(?:\s*\([^)]*\))?$",
    r"^citation style guide(?:\s+link)?(?:\s*\([^)]*\))?$",
    r"^citable levels?(?:\s*\([^)]*\))?$",
    r"^citation(?:\s+format\s+requirements|\s+standardization)?\s+rules?(?:\s*\([^)]*\))?$",
    r"^toc\b.*sorting order(?:\s*\([^)]*\))?$",
    r"^toc\b.*hiding levels?(?:\s*\([^)]*\))?$",
    r"^content profil(?:e|ing)(?:\s*\([^)]*\))?$",
    r"^references?(?:\s*\([^)]*\))?$",
)


def _heading_level(paragraph: Any) -> int | None:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    match = re.search(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _looks_like_section_heading(text: str) -> bool:
    normalized = _normalize_heading(text)
    if not normalized or len(normalized) > 120:
        return False

    # Treat only compact title-like phrases as section headings. This avoids
    # misclassifying explanatory sentences such as
    # "Citation rules stand for how the citations should appear in ELA." as a
    # new heading, which would prematurely truncate the section note block.
    return any(re.search(pattern, normalized, re.IGNORECASE) for pattern in _SECTION_HEADING_PATTERNS)


def _is_heading_paragraph(paragraph: Any) -> bool:
    if _heading_level(paragraph) is not None:
        return True
    return _looks_like_section_heading(getattr(paragraph, "text", ""))


def _iter_block_items(doc):
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield "paragraph", Paragraph(child, doc)
        elif tag == "tbl":
            yield "table", Table(child, doc)


def _get_hyperlink_target(paragraph: Any, hyperlink) -> str:
    rel_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
    if rel_id:
        try:
            return str(paragraph.part.rels[rel_id].target_ref or "").strip()
        except Exception:
            return ""

    anchor = hyperlink.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor")
    return f"#{anchor}" if anchor else ""


def _paragraph_value(paragraph: Any, rich: bool = False) -> str:
    if not rich:
        return _clean(paragraph.text)

    fragments: list[str] = []
    for child in paragraph._p.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "r":
            fragment = _format_run_element(child, paragraph)
            if fragment:
                fragments.append(fragment)
        elif tag == "hyperlink":
            hyperlink_parts = [
                _format_run_element(run_element, paragraph)
                for run_element in child.xpath('.//*[local-name()="r"]')
            ]
            hyperlink_text = "".join(part for part in hyperlink_parts if part).strip()
            if hyperlink_text:
                href = _get_hyperlink_target(paragraph, child)
                if href:
                    fragments.append(f'<a href="{html_lib.escape(href)}">{hyperlink_text}</a>')
                else:
                    fragments.append(hyperlink_text)

    return _clean_rich_text("".join(fragments))


def _paragraph_block_text(paragraph: Any, rich: bool = False) -> str:
    plain_text = _clean(paragraph.text)
    if not plain_text:
        return ""

    text = _paragraph_value(paragraph, rich=rich) if rich else plain_text
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if "list" in style_name.lower() and not re.match(r"^[\u2022\-\*]\s+", plain_text):
        return f"• {text}"
    return text


def _extract_section_block(items, start_idx: int, rich: bool = False) -> tuple[list[str], list[Any]]:
    texts: list[str] = []
    tables: list[Any] = []

    for index in range(start_idx + 1, len(items)):
        kind, block = items[index]
        if kind == "paragraph":
            paragraph = block
            if _is_heading_paragraph(paragraph):
                break
            text = _paragraph_block_text(paragraph, rich=rich)
            if text:
                texts.append(text)
        elif kind == "table":
            tables.append(block)

    return texts, tables


def _extract_citation_style_guide(items, start_idx: int) -> dict[str, Any] | None:
    texts, tables = _extract_section_block(items, start_idx, rich=True)
    rows: list[dict[str, str]] = []
    description_parts = [text for text in texts if text]

    def looks_like_non_guide_table(table) -> bool:
        if not table.rows:
            return False
        header_cells = [_clean(cell.text).lower() for cell in table.rows[0].cells]
        joined = " | ".join(header_cells)
        return any(marker in joined for marker in [
            "document title", "reference url", "content url", "issuing authority",
            "metadata element", "document location", "source name", "content category name",
            "level", "required", "definition", "toc requirements", "source of law",
        ])

    for table in tables:
        if looks_like_non_guide_table(table):
            continue
        for row in table.rows:
            rich_cells = [_cell_value(cell, rich=True) for cell in row.cells]
            plain_cells = [_clean(cell.text) for cell in row.cells]
            nonempty = [
                (plain, rich)
                for plain, rich in zip(plain_cells, rich_cells)
                if plain or re.sub(r"<[^>]+>", "", rich or "").strip()
            ]
            if not nonempty:
                continue

            normalized_first = _normalize_heading(nonempty[0][0])
            if normalized_first in {"label", "value"}:
                continue

            if len(nonempty) == 1:
                single_plain, single_rich = nonempty[0]
                single_value = _clean_rich_text(single_rich or html_lib.escape(single_plain))

                if normalized_first in {"sme checkpoint", "sme check point", "sme check-point"}:
                    if single_value:
                        description_parts.append(single_value)
                    continue

                if len(row.cells) > 1 and plain_cells[0].strip():
                    rows.append({"label": plain_cells[0].strip(), "value": ""})
                    continue

                if single_value and normalized_first not in {"label", "value"}:
                    description_parts.append(single_value)
                continue

            label = nonempty[0][0]
            value = _clean_rich_text("<br/>".join(rich for _, rich in nonempty[1:] if rich))

            if normalized_first in {"sme checkpoint", "sme check point", "sme check-point"}:
                description_value = value or html_lib.escape(label)
                if description_value:
                    description_parts.append(description_value)
                continue

            if not label and not value:
                continue
            rows.append({"label": label, "value": value})

    description = "<br/>".join(part for part in description_parts if part).strip()
    if not rows and not description:
        return None

    payload: dict[str, Any] = {}
    if description:
        payload["description"] = description
    if rows:
        payload["rows"] = rows
    return payload


def _extract_additional_sections(doc) -> dict[str, Any]:
    items = list(_iter_block_items(doc))
    extras: dict[str, Any] = {}

    for idx, (kind, block) in enumerate(items):
        if kind != "paragraph":
            continue
        if not _is_heading_paragraph(block):
            continue

        heading = _normalize_heading(getattr(block, "text", ""))
        if not heading:
            continue

        if "citation style guide link" in heading and "citationStyleGuide" not in extras:
            citation_style = _extract_citation_style_guide(items, idx)
            if citation_style:
                extras["citationStyleGuide"] = citation_style
            continue

        if "toc" in heading and "sorting order" in heading and "tocSortingOrder" not in extras:
            texts, _ = _extract_section_block(items, idx, rich=True)
            text = "\n".join(texts).strip()
            if text:
                extras["tocSortingOrder"] = text
            continue

        if "toc" in heading and "hiding levels" in heading and "tocHidingLevels" not in extras:
            texts, _ = _extract_section_block(items, idx, rich=True)
            text = "\n".join(texts).strip()
            if text:
                extras["tocHidingLevels"] = text

    return extras


def _format_run(run) -> str:
    text = run.text.replace("\xa0", " ")
    if not text:
        return ""

    font = getattr(run, "font", None)
    style = getattr(run, "style", None)
    style_name = getattr(style, "name", "") or ""
    style_id = getattr(style, "style_id", "") or ""
    style_lower = f"{style_name} {style_id}".lower()
    element_xml = getattr(getattr(run, "_element", None), "xml", "") or ""

    is_underlined = bool(
        getattr(run, "underline", False)
        or getattr(font, "underline", False)
        or "underline" in style_lower
        or re.search(r"<w:u\b", element_xml)
    )
    is_italic = bool(
        getattr(run, "italic", False)
        or getattr(font, "italic", False)
        or "emphasis" in style_lower
        or "italic" in style_lower
        or re.search(r"<w:i(?:Cs)?\b", element_xml)
    )
    is_bold = bool(
        getattr(run, "bold", False)
        or getattr(font, "bold", False)
        or "strong" in style_lower
        or "bold" in style_lower
        or re.search(r"<w:b(?:Cs)?\b", element_xml)
    )

    formatted = html_lib.escape(text).replace("\n", "<br/>")
    if is_underlined:
        formatted = f"<u>{formatted}</u>"
    if is_italic:
        formatted = f"<em>{formatted}</em>"

    if bool(getattr(font, "strike", False) or getattr(font, "double_strike", False) or re.search(r"<w:(?:strike|dstrike)\b", element_xml)):
        formatted = f"<s>{formatted}</s>"
    if is_bold:
        formatted = f"<strong>{formatted}</strong>"

    color = getattr(getattr(font, "color", None), "rgb", None)
    if not color:
        match = re.search(r'<w:color[^>]*w:val="([0-9A-Fa-f]{6,8})"', element_xml)
        if match:
            color = match.group(1)
    if color:
        formatted = f'<span style="color: #{html_lib.escape(str(color))}">{formatted}</span>'
    return formatted


def _format_run_element(run_element, paragraph) -> str:
    from docx.text.run import Run

    try:
        return _format_run(Run(run_element, paragraph))
    except Exception:
        text = "".join((node.text or "") for node in run_element.xpath('.//*[local-name()="t"]'))
        return html_lib.escape(text).replace("\n", "<br/>") if text else ""


def _cell_value(cell, rich: bool = False) -> str:
    if not rich:
        return _clean(cell.text)

    paragraphs: list[str] = []
    for paragraph in getattr(cell, "paragraphs", []):
        fragments: list[str] = []
        for child in paragraph._p.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "r":
                fragment = _format_run_element(child, paragraph)
                if fragment:
                    fragments.append(fragment)
            elif tag == "hyperlink":
                hyperlink_parts = [
                    _format_run_element(run_element, paragraph)
                    for run_element in child.xpath('.//*[local-name()="r"]')
                ]
                hyperlink_html = "".join(part for part in hyperlink_parts if part)
                if not hyperlink_html:
                    text = "".join((node.text or "") for node in child.xpath('.//*[local-name()="t"]'))
                    if text:
                        hyperlink_html = html_lib.escape(text).replace("\n", "<br/>")
                if hyperlink_html:
                    href = _get_hyperlink_target(paragraph, child)
                    if href:
                        fragments.append(f'<a href="{html_lib.escape(href)}">{hyperlink_html}</a>')
                    else:
                        fragments.append(hyperlink_html)

        if fragments:
            paragraphs.append("".join(fragments))
            continue

        plain = _clean(paragraph.text)
        if plain:
            paragraphs.append(html_lib.escape(plain))

    if not paragraphs:
        return _clean(cell.text)

    return _clean_rich_text("<br/>".join(paragraphs))


def _required_value(raw: str) -> str:
    """
    Normalise the Required cell to one of: 'Yes' | 'No' | 'Conditional' | ''
    The cell contains 'True' / 'False' in most BRDs.
    """
    val = raw.strip().lower()
    if val in ("true", "yes", "y"):
        return "Yes"
    if val in ("false", "no", "n"):
        return "No"
    if "conditional" in val or "cond" in val:
        return "Conditional"
    return ""


def _is_toc_structure_table(table) -> bool:
    """
    Return True if this table is the Document Structure / Levels table.
    Identified by a header row that mentions level, required, and definition.
    """
    if not table.rows:
        return False
    header = " ".join(c.text for c in table.rows[0].cells).lower()
    return (
        "level" in header
        and "required" in header
        and "definition" in header
    )


def _detect_col_positions(header_row) -> dict[str, int]:
    """
    Map column names to their index by scanning the header row.
    Returns a dict: { 'level': i, 'name': i, 'required': i, ... }
    Falls back to positional defaults (0-7) if a column isn't found.
    """
    keyword_map = {
        "level":           ["level"],
        "name":            ["name", "identifies level"],
        "required":        ["required", "true levels"],
        "definition":      ["definition", "level value"],
        "example":         ["example", "sample values"],
        "note":            ["note", "specific instruction"],
        "tocRequirements": ["toc requirements", "toc req", "sme checkpoint for sme"],
        "smeComments":     ["sme comments", "sme checkpoint if"],
    }
    defaults = ["level", "name", "required", "definition", "example", "note", "tocRequirements", "smeComments"]

    positions: dict[str, int] = {}
    for ci, cell in enumerate(header_row.cells):
        text = cell.text.lower().strip()
        for field, keywords in keyword_map.items():
            if field not in positions and any(kw in text for kw in keywords):
                positions[field] = ci
                break

    # Fill in positional fallbacks for anything not detected
    for pos, field in enumerate(defaults):
        if field not in positions:
            positions[field] = pos

    return positions

# Public extractor
def extract_toc(doc) -> dict:
    """
    Find the Document Structure table and extract each level row.
    Falls back to extract_toc_legacy() for paragraph-based legacy BRDs.
    """
    extras = _extract_additional_sections(doc)

    # Legacy BRDs store levels as paragraphs, not a table
    for table in doc.tables:
        if _is_toc_structure_table(table):
            break
    else:
        legacy = extract_toc_legacy(doc)
        return {**legacy, **extras}

    for table in doc.tables:
        if not _is_toc_structure_table(table):
            continue

        col = _detect_col_positions(table.rows[0])
        sections = []

        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue  # skip header

            cells = row.cells
            n = len(cells)

            def cell_text(field: str) -> str:
                idx = col.get(field)
                if idx is None or idx >= n:
                    return ""
                rich_fields = {"name", "definition", "example", "note", "tocRequirements", "smeComments"}
                return _cell_value(cells[idx], rich=field in rich_fields)

            level_cell = cell_text("level")
            level_raw = _normalize_level_value(level_cell)

            # Skip completely empty rows or long descriptive sub-headers.
            if not level_raw:
                continue
            if not re.search(r"\d", level_raw) and len(level_raw) > 5:
                continue

            definition = cell_text("definition")
            sections.append({
                "id":              level_raw,
                "level":           level_raw,
                "name":            cell_text("name"),
                "required":        _required_value(cell_text("required")),
                "definition":      definition,
                "example":         cell_text("example"),
                "note":            cell_text("note"),
                "tocRequirements": cell_text("tocRequirements"),
                "smeComments":     cell_text("smeComments"),
                # Populated for Level 0/1 whose definition is "Hardcoded – /xxx"
                "path":            _extract_path_from_definition(definition),
            })

        # Found and processed the table — return immediately
        if sections:
            return {"sections": sections, **extras}

    return {"sections": [], **extras}


# ─────────────────────────────────────────────
# Legacy extractor (paragraph-based format)
# ─────────────────────────────────────────────

_LEGACY_LEVEL_RE = re.compile(r"^level\s+(\d+)", re.IGNORECASE)
_LEGACY_FIELD_RE = re.compile(
    r"^(name|required|definition|example location|example|note|ex\.|ex:)\s*[:\-–]?\s*(.*)",
    re.IGNORECASE,
)




def extract_toc_legacy(doc) -> dict:
    """
    Legacy BRDs store levels as paragraphs instead of a table:

      Normal paragraph  → "Level N"
      List Paragraph    → "Name: ..."
                          "Required: True / False"
                          "Definition: ..."
                          "Example: ..."
                          "Note: ..."

    Collects from Heading 3 "Levels" through the next major section heading.

    FIX: The previous version broke out of the loop on *any* Heading 3
    encountered while collecting, which caused Level 0 and Level 1 to be
    dropped whenever a sub-heading appeared before their bullet lines were
    fully read.  Now only stop-keyword headings (and Heading 2+) terminate
    collection; unrecognised Heading 3s inside the Levels section are ignored.

    Output shape is identical to extract_toc().
    Each Level 0 / Level 1 section also gets a "path" key derived from its
    definition (e.g. "Hardcoded – /us" → "/us") so that content_profile_extractor
    can assemble hardcoded_path without hardcoding "/us" itself.
    """
    sections = []
    current: dict | None = None

    _stop_lower = [
        "example annotated",
        "annotated header text levels",
        "metadata",
        "references to other",
        "exceptions",
        "assumptions",
        "file delivery",
    ]

    def _flush():
        if current and current.get("level"):
            # Derive path for Level 0 / Level 1 from the definition field
            if current["level"] in ("0", "1") and not current.get("path"):
                current["path"] = _extract_path_from_definition(current["definition"])
            sections.append(current)

    collecting = False

    for p in doc.paragraphs:
        style      = p.style.name if p.style else ""
        text       = p.text.replace("\xa0", " ").strip()
        text_lower = text.lower()
        is_heading = style.startswith("Heading")

        # Start collecting at the `Levels` marker. Some legacy BRDs use a
        # plain paragraph instead of a heading for this label.
        is_levels_marker = bool(re.fullmatch(r"levels\s*:?", text_lower)) or (
            is_heading
            and "levels" in text_lower
            and "citable" not in text_lower
            and "annotated" not in text_lower
        )
        if is_levels_marker and not collecting:
            collecting = True
            continue

        if not collecting:
            continue

        # Always stop at a known stop-keyword heading (any heading level)
        if is_heading and any(s in text_lower for s in _stop_lower):
            break

        # Stop at Heading 2 (a new top-level section) — but NOT at Heading 3,
        # because sub-headings like "Annotated Header Text" live inside the
        # Levels section and should not prematurely terminate collection.
        if style.startswith("Heading 2") and collecting:
            break

        # A Heading 3 that is *not* the opening "Levels" heading and *not* a
        # stop keyword: just skip the heading line itself, keep collecting.
        if is_heading:
            continue

        # Detect "Level N" marker line
        m_level = _LEGACY_LEVEL_RE.match(text)
        if m_level:
            _flush()
            lvl = m_level.group(1)
            current = {
                "id":              lvl,
                "level":           lvl,
                "name":            "",
                "required":        "",
                "definition":      "",
                "example":         "",
                "note":            "",
                "tocRequirements": "",
                "smeComments":     "",
                "path":            "",   # populated in _flush() for L0/L1
            }
            continue

        if current is None:
            continue

        # Parse named bullet fields
        m_field = _LEGACY_FIELD_RE.match(text)
        if m_field:
            key_raw = m_field.group(1).lower().rstrip(".")
            value   = m_field.group(2).replace("\xa0", " ").strip()
            if key_raw == "name":
                current["name"] = value
            elif key_raw == "required":
                current["required"] = _required_value(value)
            elif key_raw == "definition":
                current["definition"] = value
            elif key_raw in ("example", "example location", "ex.", "ex:"):
                current["example"] = (current["example"] + "; " + value).lstrip("; ") if current["example"] else value
            elif key_raw == "note":
                current["note"] = value
        else:
            # Plain continuation — append to definition
            if text:
                if current["definition"]:
                    current["definition"] += " " + text
                else:
                    current["definition"] = text

    _flush()

    extras = _extract_additional_sections(doc)
    return {"sections": sections, **extras}