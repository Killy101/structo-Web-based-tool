"""
src/services/extractors/scope_extractor.py
Extracts the scope table from a BRD .docx file.

Supported scope formats
-----------------------
1.  Standard multi-column scope table (CFR / KR-NARK style)
      Header row with "Document Title", "Reference URL", etc.
      One document per row, with optional Issuing Authority / ASRB ID.

2.  Single-column title list under Scope heading (Alaska / Utah Code style)
      Table with 1 column, one title per row ("Title 13", "Title 17", …).

3.  Multi-column chapter-number grid under Scope heading (Hawaii style)
      Table with 2+ columns, cells are short alphanumeric codes.

4.  Legacy paragraph list under Scope heading
      Plain paragraphs, one entry per line ("Title 5", "Chapter 98", …).
      Sub-formats handled:
        4a. Simple list (Alabama, Tennessee, Arkansas, Arizona, etc.)
        4b. Title+URL pairs – next paragraph is a bare URL (Delaware style)
        4c. Inline link – "(link: <url>)" appended to title (Hawaii-style prose)

5.  Pipe-delimited inline paragraph (Oregon Revised Statutes, Utah Admin, VI)
      Single paragraph containing multiple entries separated by "|" or "│"
      e.g. "| Chapter 98 | Chapter 93 |…"  or  "Titles 152│164│331│"

6.  Department/agency name list (Wisconsin Admin Code)
      Paragraph entries that are department/agency names rather than Title/Chapter
      codes — treated as document titles as-is.

7.  Sentence-embedded scope (Texas Constitution, South Dakota, West Virginia Bar)
      Scope is described in a single sentence mentioning an article/rule number.
      Parsed via regex to extract the identifier.
"""

import re
from .base import extract_url_and_note_from_cell, extract_url_and_note_from_text, extract_url_from_cell
from .toc_extractor import _cell_value, _extract_section_block, _is_heading_paragraph, _iter_block_items, _normalize_heading


# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

_URL_COL_KEYWORDS: list[str] = [
    "reference url", "content url", "reference link",
    "parent url", "regulator url", "regulator weblink",
    "url for the title", "url for the source",
]

_TITLE_COL_KEYWORDS: list[str] = [
    "document title", "source name",
]

_HEADER_FIRST_RUN_LABELS: frozenset[str] = frozenset({
    "document title", "reference url", "content url", "reference link",
    "parent url", "regulator url", "regulator weblink",
    "issuing authority", "asrb id", "sme comments", "sme checkpoint",
    "date of ingestion", "initial evergreen", "initial/evergreen",
    "initial/ evergreen",
    "source name", "url for the title", "url for the source",
    "innodata only",
})

_BOILERPLATE_MARKERS: list[str] = [
    "innodata only", "document title as appearing", "appearing on regulator",
    "sme check", "if anything needs", "click any cell", "sorting order",
    "toc -", "*toc", "template", "example", "rolling cycle",
    "ecfr is updated", "electronic code of federal", "officially maintained",
    "smes to check", "check if weblink", "up-to-date on a rolling",
    "url for the title", "url for the source", "parent url for the source",
    "please reference", "relevant sections for us states",
    "the following titles are currently in scope",
    "titles are currently in scope",
    "currently in scope",
    "for most up to date",
    "most up to date relevant",
    "the following",
    "the scope of the source is",
    "note: ",
    "the ct statutes",
]

_URL_RE   = re.compile(r"https?://")
_ASRB_RE  = re.compile(r"\bASRB[- ]?\d+\b", re.IGNORECASE)

_INTRO_BLOB_RE = re.compile(
    r"(following\s+titles\s+are\s+currently\s+in\s+scope"
    r"|titles\s+are\s+currently\s+in\s+scope"
    r"|currently\s+in\s+scope\s*:?"
    r"|please\s+reference\b"
    r"|the\s+scope\s+of\s+the\s+source\s+is\s*:?"
    r"|chapters\s+can\s+be\s+captured\s*:?"
    r"|from\s+the\s+above\s+url\s+should\s+be\s+captured"
    r"|should\s+be\s+captured\b)",
    re.IGNORECASE,
)

_SCOPE_ENTRY_RE = re.compile(
    r"^(Title|Chapter|Rule|Part|Article|Section|Subchapter|Division|Code)\s+"
    r"[0-9A-Za-z][A-Za-z0-9\-\.]*\s*$"
    r"|^(TITLE|CHAPTER)\s+[0-9]+[A-Za-z]*\s*$",
    re.IGNORECASE,
)

_CHAPTER_ID_RE = re.compile(r"^[0-9]+[A-Za-z]*$")
_PIPE_SPLIT_RE = re.compile(r"[|│]")

_MULTI_CHAPTER_RE = re.compile(
    r"^(.+?)\s+Chapters?\s+([\d\w]+(?:\s*,\s*[\d\w]+)+)\s*$",
    re.IGNORECASE,
)

_SENTENCE_SCOPE_RE = re.compile(
    r"\b((?:Title|Chapter|Rule|Part|Article|Acticle|Section)\s+[0-9][A-Za-z0-9\-\.]*)",
    re.IGNORECASE,
)


def _expand_multi_chapter(title: str, ref_url: str = "", content_url: str = "") -> list[dict]:
    m = _MULTI_CHAPTER_RE.match(title.strip())
    if m:
        prefix   = m.group(1).strip()
        chapters = [c.strip() for c in m.group(2).split(",") if c.strip()]
        return [_make_entry(f"{prefix} Chapter {c}", ref_url, content_url) for c in chapters]
    return [_make_entry(title, ref_url, content_url)]


def _make_entry(title: str, ref_url: str = "", content_url: str = "") -> dict:
    return {
        "document_title":         title,
        "regulator_url":          ref_url,
        "content_url":            content_url,
        "content_note":           "",
        "issuing_authority":      "",
        "issuing_authority_code": "",
        "geography":              "",
        "asrb_id":                "",
        "sme_comments":           "",
        "initial_evergreen":      "",
        "date_of_ingestion":      "",
        "strikethrough":          False,
    }


def _normalise_asrb_and_sme(asrb_raw: str, sme_comments: str, fallback_asrb: str = "") -> tuple[str, str]:
    """Keep ASRB IDs in the ASRB field even when older BRDs placed them in SME comments."""
    def _dedupe(ids: list[str]) -> str:
        seen: set[str] = set()
        ordered: list[str] = []
        for value in ids:
            normalized = value.upper().replace(" ", "").replace("-", "")
            if normalized and normalized not in seen:
                seen.add(normalized)
                ordered.append(normalized)
        return ", ".join(ordered)

    asrb_id = _dedupe(_ASRB_RE.findall(asrb_raw or "")) or (fallback_asrb or "")
    cleaned_comments = (sme_comments or "").strip()

    if cleaned_comments:
        embedded_asrb = _dedupe(_ASRB_RE.findall(cleaned_comments))
        if embedded_asrb:
            asrb_id = asrb_id or embedded_asrb
            cleaned_comments = _ASRB_RE.sub("", cleaned_comments)
            cleaned_comments = re.sub(r"^[\s,;:/-]+|[\s,;:/-]+$", "", cleaned_comments)
            cleaned_comments = re.sub(r"\s{2,}", " ", cleaned_comments).strip()

    return asrb_id, cleaned_comments


def _extract_scope_sme_checkpoint(doc) -> str:
    items = list(_iter_block_items(doc))
    for idx, (kind, block) in enumerate(items):
        if kind != "paragraph":
            continue
        if not _is_heading_paragraph(block):
            continue

        heading = _normalize_heading(getattr(block, "text", ""))
        if not heading or not heading.startswith("scope"):
            continue

        texts, _ = _extract_section_block(items, idx, rich=True)
        note = "\n".join(texts).strip()
        if note:
            return note

    return ""


def _with_scope_checkpoint(result: dict, scope_sme_checkpoint: str) -> dict:
    if scope_sme_checkpoint and isinstance(result, dict) and not result.get("smeCheckpoint"):
        updated = dict(result)
        updated["smeCheckpoint"] = scope_sme_checkpoint
        return updated
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Low-level cell helpers
# ─────────────────────────────────────────────────────────────────────────────

def _cell_is_strikethrough(cell) -> bool:
    strike_runs = total_runs = 0

    for run in cell._tc.xpath('.//*[local-name()="r"]'):
        text = "".join((node.text or "") for node in run.xpath('.//*[local-name()="t"]')).strip()
        if not text:
            continue
        total_runs += 1
        has_strike = bool(
            run.xpath('./*[local-name()="rPr"]/*[local-name()="strike" or local-name()="dstrike"]')
        )
        if has_strike:
            strike_runs += 1

    if total_runs == 0:
        for para in cell.paragraphs:
            for run in para.runs:
                if not run.text.strip():
                    continue
                total_runs += 1
                if run.font.strike:
                    strike_runs += 1

    return total_runs > 0 and strike_runs == total_runs


def _row_is_strikethrough(row, preferred_col: int | None = None) -> bool:
    cells = row.cells
    if preferred_col is not None and preferred_col < len(cells) and _cell_is_strikethrough(cells[preferred_col]):
        return True
    return any(_cell_is_strikethrough(cell) for cell in cells)


def _get_first_run_text(cell) -> str:
    for para in cell.paragraphs:
        for run in para.runs:
            t = run.text.strip().lower()
            if t:
                return t
    return cell.text.strip().lower()


def _row_is_header(row) -> bool:
    if not row.cells:
        return False
    first = _get_first_run_text(row.cells[0])
    return first in _HEADER_FIRST_RUN_LABELS


# ─────────────────────────────────────────────────────────────────────────────
# Table-type detection helpers
# ─────────────────────────────────────────────────────────────────────────────

def _is_chapter_grid_table(table) -> bool:
    if not table.rows or len(table.rows) < 3:
        return False
    if len(table.rows[0].cells) < 2:
        return False
    total = matching = 0
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip().replace("\xa0", " ")
            if not text:
                continue
            total += 1
            if _CHAPTER_ID_RE.match(text):
                matching += 1
    return total > 0 and (matching / total) >= 0.60


def _find_scope_section_table(doc):
    body_els = list(doc.element.body)
    in_scope_section = False

    for el in body_els:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc)
            style = (p.style.name or "") if p.style else ""
            text_lower = p.text.strip().lower()
            if style.startswith("Heading"):
                if "scope" in text_lower:
                    in_scope_section = True
                elif in_scope_section:
                    break

        elif tag == "tbl" and in_scope_section:
            from docx.table import Table as DocxTable
            t = DocxTable(el, doc)
            ncols = len(t.columns)
            nrows = len(t.rows)

            # Single-column list
            if ncols == 1 and nrows >= 2:
                titles = [row.cells[0].text.strip() for row in t.rows if row.cells[0].text.strip()]
                real = [tt for tt in titles if not any(m in tt.lower() for m in _BOILERPLATE_MARKERS)]
                if len(real) >= 1:
                    return t

            # Multi-column chapter-number grid
            if ncols >= 2 and _is_chapter_grid_table(t):
                return t

            # Non-grid multi-col table — stop searching
            if ncols > 1:
                break

    return None


# ─────────────────────────────────────────────────────────────────────────────
# Table scoring & selection (standard scope table)
# ─────────────────────────────────────────────────────────────────────────────

def _score_table(table) -> int:
    if not table.rows:
        return -999
    row0_cells = table.rows[0].cells
    ncols      = len(row0_cells)
    nrows      = len(table.rows)
    all_header = " ".join(c.text.lower().strip() for c in row0_cells)
    score = 0
    has_url_col   = any(kw in all_header for kw in _URL_COL_KEYWORDS)
    has_title_col = any(kw in all_header for kw in _TITLE_COL_KEYWORDS)
    if has_url_col:   score += 20
    if has_title_col: score += 20
    if has_url_col and has_title_col: score += 30
    if ncols < 3:  score -= 50
    if nrows == 2: score -= 5
    elif nrows < 3: score -= 10
    url_rows = sum(
        1 for row in table.rows
        if _URL_RE.search(" ".join(c.text for c in row.cells))
    )
    score += url_rows * 5
    return score


def _find_scope_table(doc):
    best_table = None
    best_score = 0
    for table in doc.tables:
        s = _score_table(table)
        if s > best_score:
            best_score = s
            best_table = table
    return best_table


_detect_scope_table = _find_scope_table


# ─────────────────────────────────────────────────────────────────────────────
# Pre-header context detection
# ─────────────────────────────────────────────────────────────────────────────

def _detect_pre_header_context(doc, scope_table) -> dict:
    context  = {"issuing_authority": "", "asrb_id": ""}
    scope_el = scope_table._tbl
    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "tbl":
            if block is scope_el:
                break
            from docx.table import Table as DocxTable
            t        = DocxTable(block, doc)
            all_text = " ".join(c.text for row in t.rows for c in row.cells)
            if "ASRB" not in all_text or "issuing" not in all_text.lower():
                continue
            for row in t.rows:
                cells    = [c.text.strip() for c in row.cells]
                row_text = " ".join(cells).lower()
                if "issuing authority" in row_text and "asrb id" in row_text:
                    continue
                if "innodata to capture" in row_text:
                    continue
                if len(cells) >= 2:
                    auth = cells[0]
                    if auth and not any(
                        kw in auth.lower() for kw in ["issuing", "innodata", "note", "asrb", "sme"]
                    ):
                        context["issuing_authority"] = auth.strip()
                    asrb_match = _ASRB_RE.search(cells[1])
                    if asrb_match:
                        context["asrb_id"] = asrb_match.group(0)
    return context


# ─────────────────────────────────────────────────────────────────────────────
# Column-map detection
# ─────────────────────────────────────────────────────────────────────────────

def _detect_column_map(header_rows: list) -> dict:
    col_texts: dict[int, str] = {}
    for row in header_rows:
        for ci, cell in enumerate(row.cells):
            col_texts[ci] = col_texts.get(ci, "") + " " + cell.text.lower().strip()

    col_map: dict[str, int | None] = {
        "title": None, "ref_url": None, "content_url": None,
        "authority": None, "asrb_id": None, "sme": None,
        "initial_evergreen": None, "date_of_ingestion": None,
    }
    keyword_rules = [
        ("title",             ["document title", "innodata only", "title as appearing", "source name"]),
        ("ref_url",           ["reference url", "reference link", "parent url", "regulator weblink", "regulator url"]),
        ("content_url",       ["content url", "url for the title", "content link", "url for the source"]),
        ("authority",         ["issuing authority", "innodata to capture"]),
        ("initial_evergreen", ["initial/\nevergreen", "initial/ evergreen", "initial evergreen", "evergreen"]),
        ("date_of_ingestion", ["date of ingestion", "ingestion date"]),
        ("asrb_id",           ["asrb id", "asrb\n", "asrb "]),
        ("sme",               ["sme comments", "sme checkpoint"]),
    ]
    assigned: set[int] = set()
    for field, keywords in keyword_rules:
        for ci, text in sorted(col_texts.items()):
            if ci in assigned:
                continue
            if any(kw in text for kw in keywords):
                col_map[field] = ci
                assigned.add(ci)
                break
    assigned_positions = {v for v in col_map.values() if v is not None}
    for pos, field in enumerate(["title", "ref_url", "content_url", "authority", "asrb_id", "sme"]):
        if col_map[field] is None and pos not in assigned_positions:
            col_map[field] = pos
            assigned_positions.add(pos)
    return col_map


# ─────────────────────────────────────────────────────────────────────────────
# Data-start row detection
# ─────────────────────────────────────────────────────────────────────────────

def _normalise(text: str) -> str:
    plain = re.sub(r"<[^>]+>", " ", text or "")
    s = re.sub(r"[‐‑‒–—−]+", "-", plain.lower())
    s = re.sub(r"[\u2018\u2019\u201a\u201b\u2032\u2035]", "'", s)
    s = re.sub(r"[\u201c\u201d\u201e\u201f\u2033\u2036]", '"', s)
    return re.sub(r"\s+", " ", s).strip()


_NON_DATA_MARKERS: list[str] = [
    *_BOILERPLATE_MARKERS,
    "sme check-point", "sme checkpoint",
    "if anything needs be changed",
    "click any cell", "hover row", "sorting order",
    "toc - sorting order", "*toc - sorting order",
    "document title as appearing on regulator weblink",
    "url for the title under the source",
    "checkpoint", "appearing on regulator weblink",
    "please reference", "relevant sections for us states",
    "the following titles are currently in scope",
    "titles are currently in scope", "currently in scope",
    "for most up to date", "most up to date relevant sections",
    "the following", "the scope of the source is",
]


def _is_non_data_scope_row(doc_title, ref_url, content_url, sme_comments) -> bool:
    title_n = _normalise(doc_title or "")
    sme_n   = _normalise(sme_comments or "")
    has_url = bool((ref_url or "").strip() or (content_url or "").strip())
    if not title_n and not has_url:
        return True
    if any(m in title_n for m in _NON_DATA_MARKERS): return True
    # SME comments often contain phrases like "the following" or template guidance
    # even on real scope rows. Only let comment boilerplate suppress a row when the
    # row itself is otherwise blank or lacks any document URL signal.
    if not title_n and any(m in sme_n for m in _NON_DATA_MARKERS): return True
    if not has_url and any(m in sme_n for m in _NON_DATA_MARKERS): return True
    if _INTRO_BLOB_RE.search(title_n): return True
    if re.match(r"^\*", title_n.strip()): return True
    if re.search(r"\bsme\s*check[-\s]*point\b", title_n): return True
    if re.search(r"\btoc\s*[-\s]*sorting\s*order\b", title_n): return True
    if re.search(r"\bif\s+anything\s+needs\s+be\s+changed\b", title_n): return True
    if re.search(r"\bclick\s+any\s+cell\s+to\s+edit\b", title_n): return True
    if not has_url and len(title_n) > 80: return True
    return False


def _find_data_start_row(table, col_map: dict) -> int:
    title_col   = col_map.get("title", 0)
    content_col = col_map.get("content_url", 2)
    ref_col     = col_map.get("ref_url", 1)
    first_candidate = None
    for ri, row in enumerate(table.rows):
        cells = row.cells
        n     = len(cells)
        if _row_is_header(row):
            continue
        title_text   = cells[title_col].text.strip()   if title_col   is not None and title_col   < n else ""
        content_text = cells[content_col].text.strip() if content_col is not None and content_col < n else ""
        ref_text     = cells[ref_col].text.strip()     if ref_col     is not None and ref_col     < n else ""
        title_lower  = _normalise(title_text)
        if any(m in title_lower for m in _BOILERPLATE_MARKERS):
            continue
        if _INTRO_BLOB_RE.search(title_lower):
            continue
        if title_lower.strip().startswith("*"):
            continue
        has_url = bool(_URL_RE.search(content_text) or _URL_RE.search(ref_text))
        if not has_url and len(title_lower) > 80:
            continue
        if _is_non_data_scope_row(title_text, ref_text, content_text, ""):
            continue
        if first_candidate is None and (title_text or has_url):
            first_candidate = ri
        if has_url and first_candidate is None:
            return ri
    return first_candidate if first_candidate is not None else len(table.rows)


# ─────────────────────────────────────────────────────────────────────────────
# Legacy-format detection
# ─────────────────────────────────────────────────────────────────────────────

def _is_legacy_format(doc) -> bool:
    has_scope_heading = any(
        p.style and p.style.name.startswith("Heading") and "scope" in p.text.lower()
        for p in doc.paragraphs
    )
    if not has_scope_heading:
        return False
    scope_table = _find_scope_table(doc)
    if scope_table is None:
        return True
    total_rows  = len(scope_table.rows)
    header_rows = scope_table.rows[:min(4, total_rows)]
    col_map     = _detect_column_map(header_rows)
    data_start  = _find_data_start_row(scope_table, col_map)
    title_col   = col_map.get("title", 0)
    ref_col     = col_map.get("ref_url", 1)
    content_col = col_map.get("content_url", 2)
    useful_rows = 0
    for row in scope_table.rows[data_start:]:
        cells = row.cells
        n     = len(cells)
        if title_col is not None and title_col < n:
            title       = cells[title_col].text.strip()
            # Pass actual URL columns so long-title rows (>80 chars) are not
            # wrongly rejected by the "no URL and len > 80" guard in
            # _is_non_data_scope_row — which would cause the extractor to
            # fall through to extract_scope_legacy and pick up boilerplate.
            ref_url     = cells[ref_col].text.strip()     if ref_col     is not None and ref_col     < n else ""
            content_url = cells[content_col].text.strip() if content_col is not None and content_col < n else ""
            if title and not _is_non_data_scope_row(title, ref_url, content_url, ""):
                useful_rows += 1
    return useful_rows == 0


# ─────────────────────────────────────────────────────────────────────────────
# Legacy-section paragraph helpers
# ─────────────────────────────────────────────────────────────────────────────

def _legacy_paragraphs_in_section(doc, start_heading: str, stop_headings: list[str]):
    """Yield paragraphs under *start_heading* until the next *stop_heading*."""
    collecting = False
    stop_lower = [s.lower() for s in stop_headings]
    for p in doc.paragraphs:
        style      = p.style.name if p.style else ""
        is_heading = style.startswith("Heading")
        text_lower = p.text.strip().lower()
        if is_heading and start_heading.lower() in text_lower:
            collecting = True
            continue
        if collecting:
            if is_heading and any(s in text_lower for s in stop_lower):
                break
            yield p


_SCOPE_STOP_HEADINGS = [
    "How to Identify", "Document Structure", "Metadata", "References",
    "Citation", "Levels", "Structuring", "Source", "Delivery",
    "Assumptions", "Exceptions", "File", "Appendix",
]

_INTRO_PREFIXES = (
    "the following", "please reference", "please note", "note:",
    "for most up to date", "titles are currently in scope",
    "the scope of the source is", "the scope of",
    "the codes are separated", "the following chapters",
    "the following titles", "the following rules",
    "the ct statutes",
)


def _is_intro_paragraph(text_norm: str) -> bool:
    if any(text_norm.startswith(pfx) for pfx in _INTRO_PREFIXES):
        return True
    if _INTRO_BLOB_RE.search(text_norm):
        return True
    if re.match(r"^note\s*:", text_norm):
        return True
    return False


def _get_source_url(doc) -> str:
    url_re = re.compile(r"https?://\S+")
    for p in _legacy_paragraphs_in_section(
        doc, "Source", ["Scope", "How to Identify", "Document Structure"]
    ):
        m = url_re.search(p.text)
        if m:
            return m.group(0).rstrip(".,;")
    for p in _legacy_paragraphs_in_section(doc, "Scope", _SCOPE_STOP_HEADINGS):
        text = p.text.strip()
        if _URL_RE.match(text):
            return text.rstrip(".,;")
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# Format 5: Pipe-delimited inline paragraph parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pipe_delimited(text: str, ref_url: str = "") -> list[dict]:
    prefix_match = re.match(r"^(Title[s]?|Chapter[s]?|Rule[s]?)\s*[\u2013\-]?\s*", text, re.IGNORECASE)
    prefix_word = ""
    if prefix_match:
        raw_word    = prefix_match.group(1).rstrip("sS").strip()
        prefix_word = raw_word.capitalize()
        text        = text[prefix_match.end():]

    parts = _PIPE_SPLIT_RE.split(text)
    entries = []
    seen: set[str] = set()
    for part in parts:
        title = part.strip().replace("\xa0", " ")
        if not title:
            continue
        if prefix_word and re.match(r"^[0-9]+[A-Za-z]*$", title):
            title = f"{prefix_word} {title}"
        if any(m in title.lower() for m in _BOILERPLATE_MARKERS):
            continue
        if title in seen:
            continue
        seen.add(title)
        entries.append(_make_entry(title, ref_url))
    return entries


def _is_pipe_delimited(text: str) -> bool:
    return len(_PIPE_SPLIT_RE.findall(text)) >= 2


# ─────────────────────────────────────────────────────────────────────────────
# Format 7: Sentence-embedded scope parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_sentence_scope(text: str, ref_url: str = "") -> list[dict]:
    matches = _SENTENCE_SCOPE_RE.findall(text)
    entries = []
    seen: set[str] = set()
    for m in matches:
        title = m.strip()
        if title and title not in seen:
            seen.add(title)
            entries.append(_make_entry(title, ref_url))
    return entries


# ─────────────────────────────────────────────────────────────────────────────
# Table-based extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_scope_single_col_table(table, ref_url: str = "") -> dict:
    in_scope = []
    seen: set[str] = set()
    for row in table.rows:
        title = row.cells[0].text.strip().replace("\xa0", " ")
        if not title:
            continue
        if any(m in title.lower() for m in _BOILERPLATE_MARKERS):
            continue
        if _INTRO_BLOB_RE.search(title.lower()):
            continue
        if title in seen:
            continue
        seen.add(title)
        in_scope.append(_make_entry(title, ref_url))
    return {
        "in_scope": in_scope, "out_of_scope": [],
        "summary": f"Scope covers {len(in_scope)} active documents (single-column table format).",
    }


def _extract_scope_chapter_grid(table, ref_url: str = "") -> dict:
    in_scope = []
    seen: set[str] = set()
    for row in table.rows:
        for cell in row.cells:
            title = cell.text.strip().replace("\xa0", " ")
            if not title or title in seen:
                continue
            if any(m in title.lower() for m in _BOILERPLATE_MARKERS):
                continue
            seen.add(title)
            in_scope.append(_make_entry(title, ref_url))
    return {
        "in_scope": in_scope, "out_of_scope": [],
        "summary": f"Scope covers {len(in_scope)} chapters/titles (grid format).",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Legacy paragraph extractor (formats 4, 5, 6, 7)
# ─────────────────────────────────────────────────────────────────────────────

def extract_scope_legacy(doc) -> dict:
    ref_url         = _get_source_url(doc)
    _inline_link_re = re.compile(r"\s*\(link:\s*(https?://\S+?)\)\s*$", re.IGNORECASE)
    url_re          = re.compile(r"https?://\S+")

    def _is_url(text: str) -> bool:
        return bool(url_re.match(text))

    def _extract_inline_link(text: str) -> tuple[str, str]:
        m = _inline_link_re.search(text)
        if m:
            return _inline_link_re.sub("", text).strip(), m.group(1).rstrip(".,;)")
        return text, ""

    raw_paras: list[str] = []
    for p in _legacy_paragraphs_in_section(doc, "Scope", _SCOPE_STOP_HEADINGS):
        text = p.text.replace("\xa0", " ").strip()
        if not text:
            continue
        text_norm = _normalise(text)
        if _is_intro_paragraph(text_norm):
            if _is_pipe_delimited(text) or _SENTENCE_SCOPE_RE.search(text):
                raw_paras.append(text)
            continue
        raw_paras.append(text)

    in_scope: list[dict] = []
    seen: set[str]       = set()

    i = 0
    while i < len(raw_paras):
        text = raw_paras[i]

        if _is_url(text):
            i += 1
            continue

        if _is_pipe_delimited(text):
            for entry in _parse_pipe_delimited(text, ref_url):
                if entry["document_title"] not in seen:
                    seen.add(entry["document_title"])
                    in_scope.append(entry)
            i += 1
            continue

        title, content_url = _extract_inline_link(text)

        if not content_url:
            if i + 1 < len(raw_paras) and _is_url(raw_paras[i + 1]):
                content_url = raw_paras[i + 1].rstrip(".,;")
                i += 2
            else:
                text_norm = _normalise(text)
                is_simple = bool(
                    _SCOPE_ENTRY_RE.match(text.strip()) or
                    (len(text.split()) <= 4 and not _is_intro_paragraph(text_norm))
                )
                if not is_simple and _SENTENCE_SCOPE_RE.search(text):
                    for entry in _parse_sentence_scope(text, ref_url):
                        if entry["document_title"] not in seen:
                            seen.add(entry["document_title"])
                            in_scope.append(entry)
                    i += 1
                    continue
                i += 1
        else:
            i += 1

        title_key = re.sub(r"[\u2018\u2019\u201a\u201b\u2032\u2035\u0060\u00b4]", "'", title)
        if title_key in seen:
            continue
        expanded = _expand_multi_chapter(title, ref_url, content_url)
        for entry in expanded:
            ekey = re.sub(r"[\u2018\u2019\u201a\u201b\u2032\u2035\u0060\u00b4]", "'", entry["document_title"])
            if ekey not in seen:
                seen.add(ekey)
                in_scope.append(entry)

    if not in_scope and ref_url:
        in_scope.append(_make_entry("All chapters (see source URL)", ref_url))

    return {
        "in_scope": in_scope, "out_of_scope": [],
        "summary": f"Scope covers {len(in_scope)} active documents (legacy paragraph format).",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Public extractor
# ─────────────────────────────────────────────────────────────────────────────

def extract_scope(doc) -> dict:
    """
    Extract scope from *doc* (python-docx Document).

    Priority:
      1. Chapter-number grid table under Scope heading (Hawaii-style)
      2. Single-column table under Scope heading (Alaska/Oklahoma/Utah Code)
      3. Legacy paragraph extractor (formats 4–7)
      4. Standard scored multi-column scope table
    """
    # ── Diagnostic logging ────────────────────────────────────────────────────
    heading_styles = [
        (p.style.name if p.style else "?", p.text.strip()[:80])
        for p in doc.paragraphs
        if (p.style and p.style.name.startswith("Heading")) or "scope" in p.text.lower()
    ]
    table_summaries = []
    for i, t in enumerate(doc.tables):
        row0_text = " | ".join(c.text.strip()[:30] for c in t.rows[0].cells) if t.rows else "(empty)"
        score = _score_table(t)
        table_summaries.append(f"  table[{i}] cols={len(t.columns)} rows={len(t.rows)} score={score} row0=[{row0_text}]")
    print(f"[DEBUG scope_extractor] headings/scope paragraphs: {heading_styles}")
    print(f"[DEBUG scope_extractor] {len(doc.tables)} tables in document:")
    for s in table_summaries:
        print(s)
    # ──────────────────────────────────────────────────────────────────────────

    scope_sme_checkpoint = _extract_scope_sme_checkpoint(doc)

    section_table = _find_scope_section_table(doc)
    if section_table is not None:
        ref_url = _get_source_url(doc)
        if len(section_table.columns) >= 2:
            return _with_scope_checkpoint(_extract_scope_chapter_grid(section_table, ref_url), scope_sme_checkpoint)
        return _with_scope_checkpoint(_extract_scope_single_col_table(section_table, ref_url), scope_sme_checkpoint)

    if _is_legacy_format(doc):
        print("[DEBUG scope_extractor] Using legacy format extractor")
        return _with_scope_checkpoint(extract_scope_legacy(doc), scope_sme_checkpoint)

    scope_table = _find_scope_table(doc)
    if scope_table is None:
        print("[DEBUG scope_extractor] No scope table found — returning empty scope")
        return _with_scope_checkpoint({"in_scope": [], "out_of_scope": [], "summary": "Scope table not found."}, scope_sme_checkpoint)

    pre_ctx     = _detect_pre_header_context(doc, scope_table)
    global_auth = pre_ctx["issuing_authority"]
    global_asrb = pre_ctx["asrb_id"]
    total_rows  = len(scope_table.rows)
    header_rows = scope_table.rows[:min(4, total_rows)]
    col_map     = _detect_column_map(header_rows)
    data_start  = _find_data_start_row(scope_table, col_map)

    def safe_text(row_cells, col_idx, rich: bool = False):
        if col_idx is None or col_idx >= len(row_cells):
            return ""
        cell = row_cells[col_idx]
        if rich:
            try:
                rich_text = _cell_value(cell, rich=True)
                if rich_text:
                    rich_text = re.sub(r"</?(?:s|strike|del)>", "", rich_text, flags=re.IGNORECASE)
                    return rich_text.strip()
            except Exception:
                pass
        return cell.text.strip().replace("\xa0", " ")

    def safe_url_and_note(row_cells, col_idx):
        if col_idx is None or col_idx >= len(row_cells):
            return "", ""
        return extract_url_and_note_from_cell(row_cells[col_idx])

    _auth_re = re.compile(r"^(.*?)\s*\(([^)]+)\)\s*/\s*(.+)$")

    def parse_authority(raw: str) -> tuple[str, str, str]:
        m = _auth_re.match(raw)
        if m:
            return m.group(1).strip(), m.group(2).strip(), m.group(3).strip()
        return raw, "", ""

    all_entries: list[dict] = []
    for row in scope_table.rows[data_start:]:
        cells       = row.cells
        doc_title            = safe_text(cells, col_map["title"], rich=True)
        ref_url, _           = safe_url_and_note(cells, col_map["ref_url"])
        content_url, content_note = safe_url_and_note(cells, col_map["content_url"])
        issuing_raw          = safe_text(cells, col_map["authority"]).replace("\n", " ") or global_auth
        asrb_raw             = safe_text(cells, col_map["asrb_id"])
        sme_comments_raw     = safe_text(cells, col_map["sme"], rich=True)
        asrb_id, sme_comments = _normalise_asrb_and_sme(asrb_raw, sme_comments_raw, global_asrb)
        initial_ev           = safe_text(cells, col_map.get("initial_evergreen"))
        date_ing             = safe_text(cells, col_map.get("date_of_ingestion"))

        if _is_non_data_scope_row(doc_title, ref_url, content_url, sme_comments):
            continue

        auth_name, auth_code, geography = parse_authority(issuing_raw)
        is_struck = _row_is_strikethrough(row, col_map.get("title"))
        all_entries.append({
            "document_title":         doc_title,
            "regulator_url":          ref_url,
            "content_url":            content_url,
            "content_note":           content_note,
            "issuing_authority":      auth_name,
            "issuing_authority_code": auth_code,
            "geography":              geography,
            "asrb_id":                asrb_id,
            "sme_comments":           sme_comments,
            "initial_evergreen":      initial_ev,
            "date_of_ingestion":      date_ing,
            "strikethrough":          is_struck,
        })

    active = [e for e in all_entries if not e["strikethrough"]]
    struck = [e for e in all_entries if     e["strikethrough"]]
    return _with_scope_checkpoint({
        "in_scope": active, "out_of_scope": struck,
        "summary": f"Scope covers {len(active)} active and {len(struck)} struck-through documents.",
    }, scope_sme_checkpoint)


# ─────────────────────────────────────────────────────────────────────────────
# MHTML / legacy .doc parser
# ─────────────────────────────────────────────────────────────────────────────

def _extract_scope_from_mhtml(path: str) -> dict:
    import quopri

    with open(path, "rb") as f:
        raw = f.read()

    try:
        decoded = quopri.decodestring(raw).decode("utf-8", errors="replace")
    except Exception:
        decoded = raw.decode("utf-8", errors="replace")

    def strip_tags(s: str) -> str:
        s = re.sub(r"<[^>]+>", " ", s)
        for ent, rep in [("&amp;", "&"), ("&nbsp;", " "), ("&lt;", "<"),
                         ("&gt;", ">"), ("&#39;", "'"), ("&quot;", '"')]:
            s = s.replace(ent, rep)
        s = re.sub(r"&#[0-9]+;", "", s)
        return re.sub(r"\s+", " ", s).strip()

    def extract_scope_checkpoint_from_html(html_text: str) -> str:
        section_match = re.search(r"scope(.*?)(?:document structure|citation|toc|</body>)", html_text, re.IGNORECASE | re.DOTALL)
        if not section_match:
            return ""
        segment = section_match.group(1)
        note_match = re.search(r"sme\s*checkpoint(.*?)(?:<table|document title)", segment, re.IGNORECASE | re.DOTALL)
        if not note_match:
            return ""
        note = strip_tags(note_match.group(1))
        return re.sub(r"^[:\-\s]+", "", note).strip()

    scope_sme_checkpoint = extract_scope_checkpoint_from_html(decoded)
    tables = re.findall(r"<table[^>]*>(.*?)</table>", decoded, re.DOTALL | re.IGNORECASE)

    for table_html in tables:
        rows = re.findall(r"<tr[^>]*>(.*?)</tr>", table_html, re.DOTALL | re.IGNORECASE)
        if not rows:
            continue
        header_text = strip_tags(rows[0]).lower()
        if "document title" not in header_text or "reference url" not in header_text:
            continue

        header_cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", rows[0], re.DOTALL | re.IGNORECASE)
        header_texts = [strip_tags(c).lower() for c in header_cells]

        def find_col(keywords):
            for ki, kt in enumerate(header_texts):
                if any(kw in kt for kw in keywords):
                    return ki
            return None

        title_col   = find_col(["document title"])
        ref_col     = find_col(["reference url", "parent url"])
        content_col = find_col(["content url", "url for the title"])
        auth_col    = find_col(["issuing authority"])
        asrb_col    = find_col(["asrb id"])
        sme_col     = find_col(["sme comments", "sme checkpoint"])

        in_scope: list[dict] = []
        out_of_scope: list[dict] = []
        seen: set[str] = set()

        for row_html in rows[1:]:
            cells_html = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.DOTALL | re.IGNORECASE)
            cells = [strip_tags(c) for c in cells_html]

            def gcell(idx):
                if idx is None or idx >= len(cells): return ""
                return cells[idx].strip()

            doc_title = gcell(title_col)
            if not doc_title:
                continue

            title_low = doc_title.lower()
            if any(m in title_low for m in _BOILERPLATE_MARKERS):
                continue
            if _row_is_header_text(title_low):
                continue

            ref_url, _ = extract_url_and_note_from_text(gcell(ref_col))
            content_url, content_note = extract_url_and_note_from_text(gcell(content_col))
            ref_url = re.split(r"\s+NOTE\s*:", ref_url)[0].strip()
            content_url = re.split(r"\s+NOTE\s*:", content_url)[0].strip()
            issuing_raw = gcell(auth_col)
            asrb_raw    = gcell(asrb_col)
            sme_raw     = gcell(sme_col)
            asrb_id, sme_comments = _normalise_asrb_and_sme(asrb_raw, sme_raw)

            if _is_non_data_scope_row(doc_title, ref_url, content_url, sme_comments):
                continue
            if doc_title in seen:
                continue
            seen.add(doc_title)

            _auth_re2 = re.compile(r"^(.*?)\s*\(([^)]+)\)\s*/\s*(.+)$")
            m = _auth_re2.match(issuing_raw or "")
            auth_name = m.group(1).strip() if m else issuing_raw
            auth_code = m.group(2).strip() if m else ""
            geography = m.group(3).strip() if m else ""

            entry = {
                "document_title":         doc_title,
                "regulator_url":          ref_url,
                "content_url":            content_url,
                "content_note":           content_note,
                "issuing_authority":      auth_name,
                "issuing_authority_code": auth_code,
                "geography":              geography,
                "asrb_id":                asrb_id,
                "sme_comments":           sme_comments,
                "initial_evergreen":      "",
                "date_of_ingestion":      "",
                "strikethrough":          False,
            }

            row_is_struck = bool(
                re.search(r"<(?:s|strike|del)\b", row_html, re.IGNORECASE)
                or re.search(r"text-decoration\s*:\s*line-through", row_html, re.IGNORECASE)
            )
            if row_is_struck:
                entry["strikethrough"] = True
                out_of_scope.append(entry)
            else:
                in_scope.append(entry)

        return {
            "in_scope":     in_scope,
            "out_of_scope": out_of_scope,
            "summary":      f"Scope covers {len(in_scope)} active and {len(out_of_scope)} struck-through documents (MHTML format).",
            **({"smeCheckpoint": scope_sme_checkpoint} if scope_sme_checkpoint else {}),
        }

    return {
        "in_scope": [],
        "out_of_scope": [],
        "summary": "Scope table not found in MHTML.",
        **({"smeCheckpoint": scope_sme_checkpoint} if scope_sme_checkpoint else {}),
    }


def _row_is_header_text(text_lower: str) -> bool:
    return any(lbl in text_lower for lbl in [
        "issuing authority", "innodata to capture", "asrb id",
        "sme check", "document title innodata",
    ])


def _is_mhtml_doc(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            header = f.read(256)
        return b"MIME-Version" in header or b"multipart/related" in header or b"Exported From Confluence" in header
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# File-level entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_scope_from_file(path: str) -> dict:
    import docx as _docx

    if path.lower().endswith(".doc") and _is_mhtml_doc(path):
        return _extract_scope_from_mhtml(path)

    doc = _docx.Document(path)
    return extract_scope(doc)