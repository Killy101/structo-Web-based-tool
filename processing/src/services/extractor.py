"""
src/services/extractor.py
Orchestrates all individual BRD extractors into one combined result.

extract_all() now delegates to brd_data.extract_brd() — all cleaning,
merging, and normalization happens there.  This file retains:
  - extract_text()          file → raw text (PDF / DOCX / DOC)
  - extract_all()           .docx path → combined result dict (via brd_data)
  - extract_all_sections()  async entry point for the FastAPI router
  - _fallback_from_text()   heuristic extraction when only raw text is available
"""

import re
import os
import html
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from zipfile import BadZipFile

import docx2txt
import fitz  # PyMuPDF
from docx import Document
from lxml import html as lxml_html


# ─────────────────────────────────────────────────────────────────────────────
# Public: file-type text extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_path: str, suffix: str) -> str:
    """Extract raw text from PDF, DOCX, or legacy DOC files."""
    normalized_suffix = suffix.lower()
    if normalized_suffix == ".pdf":
        return _extract_pdf(file_path)
    if normalized_suffix == ".docx":
        return _extract_docx(file_path)
    if normalized_suffix == ".doc":
        return _extract_doc(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    try:
        pages: list[str] = []
        for page in doc:
            pages.append(str(page.get_text("text")))
        return "\n\n".join(pages)
    finally:
        doc.close()


def _extract_docx(path: str) -> str:
    try:
        return docx2txt.process(path)
    except BadZipFile as exc:
        raise ValueError("Invalid .docx file content (not a valid DOCX/ZIP package).") from exc


def _extract_doc(path: str) -> str:
    """Extract text from legacy .doc by converting to a temporary .docx."""
    if _is_mhtml_doc(path):
        mhtml_text = _extract_mhtml_doc_text(path)
        if mhtml_text.strip():
            return mhtml_text

    # Use a unique path without pre-creating the file, same reason as convert_doc_to_docx.
    temp_docx_path = os.path.join(
        tempfile.gettempdir(), f"brd_extract_{uuid.uuid4().hex}.docx"
    )
    temp_docx = Path(temp_docx_path)
    try:
        if _convert_doc_to_docx_with_word(path, str(temp_docx)):
            return _extract_docx(str(temp_docx))
        if _convert_doc_to_docx_with_soffice(path, str(temp_docx)):
            return _extract_docx(str(temp_docx))
        raise ValueError(
            "Failed to read legacy .doc file. Install pywin32 (Windows + Word) "
            "or LibreOffice ('soffice' in PATH)."
        )
    finally:
        try:
            temp_docx.unlink(missing_ok=True)
        except Exception:
            pass


def _is_mhtml_doc(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            header = f.read(4096)
        lowered = header.lower()
        return (
            b"mime-version" in lowered
            or b"multipart/related" in lowered
            or b"exported from confluence" in lowered
            or b"content-location:" in lowered
            or b"<html" in lowered
        )
    except Exception:
        return False


def _extract_html_from_mhtml_doc(path: str) -> str:
    import quopri
    from email import policy
    from email.parser import BytesParser

    with open(path, "rb") as f:
        raw = f.read()

    try:
        message = BytesParser(policy=policy.default).parsebytes(raw)
        if message.is_multipart():
            html_parts: list[str] = []
            for part in message.walk():
                if part.get_content_type() != "text/html":
                    continue
                payload = part.get_payload(decode=True) or b""
                charset = part.get_content_charset() or "utf-8"
                html_parts.append(payload.decode(charset, errors="replace"))
            if html_parts:
                return "\n".join(html_parts)
    except Exception:
        pass

    try:
        decoded = quopri.decodestring(raw).decode("utf-8", errors="replace")
    except Exception:
        decoded = raw.decode("utf-8", errors="replace")

    match = re.search(r"(<html\b.*?</html>)", decoded, flags=re.IGNORECASE | re.DOTALL)
    return match.group(1) if match else decoded


def _normalize_html_text(fragment: str) -> str:
    text = re.sub(r"<[^>]+>", " ", fragment)
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _append_html_block_to_docx(node, document) -> None:
    tag = getattr(node, "tag", None)
    if not isinstance(tag, str):
        return

    tag = tag.lower()
    if tag in {"script", "style", "meta", "link", "head"}:
        return

    if tag in {"body", "div", "section", "article", "main", "header", "footer"}:
        for child in node:
            _append_html_block_to_docx(child, document)
        return

    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        text = _normalize_html_text(" ".join(node.itertext()))
        if text:
            document.add_heading(text, level=min(int(tag[1]), 6))
        return

    if tag in {"p", "li"}:
        text = _normalize_html_text(" ".join(node.itertext()))
        if text:
            document.add_paragraph(text)
        return

    if tag in {"ul", "ol"}:
        for child in node:
            _append_html_block_to_docx(child, document)
        return

    if tag == "table":
        rows: list[list[str]] = []
        max_cols = 0
        for row in node.xpath(".//tr"):
            cells = [
                _normalize_html_text(" ".join(cell.itertext()))
                for cell in row.xpath("./th|./td")
            ]
            if any(cells):
                rows.append(cells)
                max_cols = max(max_cols, len(cells))

        if rows and max_cols:
            table = document.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            for row_index, row_values in enumerate(rows):
                for col_index, value in enumerate(row_values):
                    table.cell(row_index, col_index).text = value
            document.add_paragraph("")
        return

    text = _normalize_html_text(" ".join(node.itertext()))
    if text and tag not in {"span", "strong", "em", "b", "i", "a", "br"}:
        document.add_paragraph(text)


def _convert_mhtml_doc_to_docx(src_path: str, dst_docx_path: str) -> bool:
    try:
        html_source = _extract_html_from_mhtml_doc(src_path)
        if not html_source.strip():
            return False

        root = lxml_html.fromstring(html_source)
        body = root.find("body")
        if body is None:
            body = root

        document = Document()
        before_count = len(document.paragraphs) + len(document.tables)
        for child in body:
            _append_html_block_to_docx(child, document)

        after_count = len(document.paragraphs) + len(document.tables)
        if after_count == before_count:
            text = _normalize_html_text(" ".join(root.itertext()))
            if not text:
                return False
            document.add_paragraph(text)

        document.save(dst_docx_path)
        return Path(dst_docx_path).exists() and os.path.getsize(dst_docx_path) > 0
    except Exception as exc:
        print(f"[WARN convert_doc_to_docx] HTML/MHTML conversion failed: {exc}")
        return False


def _extract_mhtml_doc_text(path: str) -> str:
    decoded = _extract_html_from_mhtml_doc(path)

    lines: list[str] = []

    # Preserve table rows as pipe-delimited lines so fallback parser can recover scope/citation rows.
    for table_html in re.findall(r"<table[^>]*>(.*?)</table>", decoded, flags=re.IGNORECASE | re.DOTALL):
        for row_html in re.findall(r"<tr[^>]*>(.*?)</tr>", table_html, flags=re.IGNORECASE | re.DOTALL):
            cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, flags=re.IGNORECASE | re.DOTALL)
            cleaned_cells = [_normalize_html_text(cell) for cell in cells]
            cleaned_cells = [cell for cell in cleaned_cells if cell]
            if cleaned_cells:
                lines.append(" | ".join(cleaned_cells))

    # Keep headings/paragraphs/list items as regular text lines.
    body_chunks = re.findall(r"<(?:h[1-6]|p|li)[^>]*>(.*?)</(?:h[1-6]|p|li)>", decoded, flags=re.IGNORECASE | re.DOTALL)
    for chunk in body_chunks:
        cleaned = _normalize_html_text(chunk)
        if cleaned:
            lines.append(cleaned)

    return "\n".join(lines)


def convert_doc_to_docx(src_path: str) -> str | None:
    """
    Convert a legacy .doc file to a temporary .docx file.

    Attempts (in order):
      0. Pure-Python HTML/MHTML synthesis for Confluence/Word-exported `.doc`.
      1. Pure-Python: the file may already be OOXML (ZIP) saved with a .doc
         extension — python-docx can open it directly without any extra tools.
      2. Word COM via pywin32 (Windows + Microsoft Word).
      3. LibreOffice soffice command-line (any platform).

    Returns the path of the converted .docx file on success (the caller is
    responsible for deleting it), or None if all methods fail.
    """
    if _is_mhtml_doc(src_path):
        temp_docx_path = os.path.join(
            tempfile.gettempdir(), f"brd_html_{uuid.uuid4().hex}.docx"
        )
        if _convert_mhtml_doc_to_docx(src_path, temp_docx_path):
            print(f"[DEBUG convert_doc_to_docx] Converted HTML/MHTML .doc via python-docx: {temp_docx_path}")
            return temp_docx_path
        Path(temp_docx_path).unlink(missing_ok=True)

    # ── Method 1: file is already OOXML (.docx) saved with a .doc extension ──
    # Many "doc" files from modern Word are actually ZIP/OOXML — python-docx
    # can open them directly.  No extra tools required.
    try:
        with zipfile.ZipFile(src_path, "r"):
            pass  # confirms it's a valid ZIP (OOXML)
        from docx import Document
        Document(src_path)  # confirms python-docx can parse it
        # Copy to a .docx temp path so downstream code sees the right extension
        temp_fd, temp_docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(temp_fd)
        shutil.copy2(src_path, temp_docx_path)
        print(f"[DEBUG convert_doc_to_docx] File is OOXML — copied as .docx (no converter needed)")
        return temp_docx_path
    except Exception:
        pass  # not OOXML; fall through to external converters

    # ── Methods 2 & 3: true binary .doc — needs Word COM or LibreOffice ───────
    # Generate a unique path WITHOUT pre-creating the file.
    # mkstemp would create an empty file that can prevent Word's SaveAs from
    # overwriting it on some Windows configurations.
    temp_docx_path = os.path.join(
        tempfile.gettempdir(), f"brd_convert_{uuid.uuid4().hex}.docx"
    )
    temp_docx = Path(temp_docx_path)
    try:
        if _convert_doc_to_docx_with_word(src_path, temp_docx_path):
            print(f"[DEBUG convert_doc_to_docx] Converted via Word COM: {temp_docx_path}")
            return temp_docx_path
        if _convert_doc_to_docx_with_soffice(src_path, temp_docx_path):
            print(f"[DEBUG convert_doc_to_docx] Converted via LibreOffice: {temp_docx_path}")
            return temp_docx_path
    except Exception as exc:
        print(f"[WARN convert_doc_to_docx] Conversion error: {exc}")
    temp_docx.unlink(missing_ok=True)
    print("[WARN convert_doc_to_docx] All conversion methods failed. "
          "Install pywin32 (Windows+Word) or LibreOffice to handle binary .doc files.")
    return None


def _convert_doc_to_docx_with_word(src_path: str, dst_docx_path: str) -> bool:
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return False

    word = None
    document = None
    initialized = False
    try:
        pythoncom.CoInitialize()
        initialized = True
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(os.path.abspath(src_path), ReadOnly=True)
        document.SaveAs(os.path.abspath(dst_docx_path), FileFormat=16)
        document.Close(False)
        document = None
        # Verify the file was actually written with real content (not 0 bytes)
        return Path(dst_docx_path).exists() and os.path.getsize(dst_docx_path) > 100
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if initialized:
            pythoncom.CoUninitialize()


def _convert_doc_to_docx_with_soffice(src_path: str, dst_docx_path: str) -> bool:
    source = Path(src_path)
    target = Path(dst_docx_path)
    with tempfile.TemporaryDirectory() as out_dir:
        cmd = ["soffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, str(source)]
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False
        if result.returncode != 0:
            return False
        converted = Path(out_dir) / f"{source.stem}.docx"
        if not converted.exists():
            return False
        target.write_bytes(converted.read_bytes())
        return True


# ─────────────────────────────────────────────────────────────────────────────
# Public: .docx → structured dict  (delegates to brd_data)
# ─────────────────────────────────────────────────────────────────────────────

def extract_all(docx_path: str, brd_id: str | None = None) -> dict:
    """
    Run all extractors on a .docx file and return the combined result dict.

    Delegates to brd_data.extract_brd() for all extraction, cleaning, and
    merging.  brd_to_metajson_input() converts the result back to the same
    dict shape this function has always returned, so all callers continue to
    work without changes.

    Parameters
    ----------
    docx_path : path to the .docx file on disk
    brd_id    : when provided, cell images are also extracted and included
                in the returned dict under "cell_images"
    """
    print(f"[DEBUG extract_all] Starting extraction for {docx_path}, brd_id={brd_id}")

    from .brd_data import extract_brd, brd_to_metajson_input

    brd    = extract_brd(docx_path, brd_id=brd_id)
    result = brd_to_metajson_input(brd)

    print(f"[DEBUG extract_all] Done. levels={len(brd.levels)}, "
          f"scope={len(brd.scope_entries)}, images={len(brd.cell_images)}")
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Async entry point for process.py router
# ─────────────────────────────────────────────────────────────────────────────

async def extract_all_sections(
    docx_path_or_text: str,
    format: str = "new",
    brd_id: str | None = None,
) -> dict:
    """
    Async entry point called by the FastAPI router.

    - If given a .docx file path → full structured extraction via brd_data.
    - If given raw text (PDF / scraper output) → heuristic fallback.
      Cell images are not available from raw text; cell_images will be [].
    """
    arg = docx_path_or_text.strip()
    is_docx_path = (
        len(arg) < 512
        and arg.endswith(".docx")
        and "\n" not in arg
        and Path(arg).exists()
    )

    if is_docx_path:
        print(f"[DEBUG extract_all_sections] Calling extract_all with brd_id={brd_id}")
        return extract_all(arg, brd_id=brd_id)

    print("[DEBUG extract_all_sections] Using fallback from text")
    return _fallback_from_text(docx_path_or_text, format)


# ─────────────────────────────────────────────────────────────────────────────
# Raw-text fallback  (PDF / .doc scraper output — no python-docx available)
# ─────────────────────────────────────────────────────────────────────────────

def _fallback_from_text(text: str, format: str) -> dict:
    """
    Heuristic extraction from raw text.
    Used when the input is not a .docx path (e.g. PDF-scraped text).
    Cell images are never available here — cell_images is always [].
    """
    import json as _json

    lines   = [line.strip() for line in text.splitlines() if line.strip()]
    joined  = "\n".join(lines)
    lowered = joined.lower()
    url_pattern = re.compile(r"https?://[^\s\)\]]+")

    # ── TOC ──────────────────────────────────────────────────────────────────
    toc_sections: list[dict] = []
    in_toc = False
    for line in lines:
        if re.search(r"table of contents", line, re.IGNORECASE):
            in_toc = True
            continue
        if in_toc:
            if len(line) > 160:
                in_toc = False
                continue
            num_match = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)", line)
            if num_match:
                number = num_match.group(1)
                title  = num_match.group(2).strip()
                level  = number.count(".") + 1
            else:
                number = str(len(toc_sections) + 1)
                title  = line
                level  = 1
            toc_sections.append({"number": number, "title": title, "page": None, "level": level})

    # ── SCOPE ─────────────────────────────────────────────────────────────────
    in_scope: list[dict] = []
    pipe_rows = [l for l in lines if l.count("|") >= 3]
    if pipe_rows:
        for row in pipe_rows:
            parts = [p.strip() for p in row.split("|")]
            if len(parts) < 4 or "document title" in parts[0].lower():
                continue
            issuing_auth = parts[3] if len(parts) > 3 else ""
            auth_match = re.match(r"^(.*?)\s*\(([^)]+)\)\s*/\s*(.+)$", issuing_auth)
            if auth_match:
                auth_name = auth_match.group(1).strip()
                auth_code = auth_match.group(2).strip()
                geography = auth_match.group(3).strip()
            else:
                auth_name, auth_code, geography = issuing_auth, "", ""
            in_scope.append({
                "document_title":         parts[0],
                "regulator_url":          parts[1] if "http" in parts[1] else "",
                "content_url":            parts[2] if "http" in parts[2] else "",
                "issuing_authority":      auth_name,
                "issuing_authority_code": auth_code,
                "geography":              geography,
                "asrb_id":                parts[4] if len(parts) > 4 else "",
                "sme_comments":           parts[5] if len(parts) > 5 else "",
                "strikethrough":          False,
            })
    else:
        for url in url_pattern.findall(joined):
            if "legislation.gov.au" in url and "/latest/text" in url:
                in_scope.append({
                    "document_title":         url,
                    "regulator_url":          "https://www.legislation.gov.au",
                    "content_url":            url,
                    "issuing_authority":      "",
                    "issuing_authority_code": "",
                    "geography":              "",
                    "asrb_id":                "",
                    "sme_comments":           "",
                    "strikethrough":          False,
                })

    # ── METADATA ──────────────────────────────────────────────────────────────
    date_match    = re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", joined)
    version_match = re.search(r"\bv(?:ersion)?\s*[:\-]?\s*([\d.]+)\b", lowered)

    metadata: dict = {
        "document_title":         lines[0] if lines else "",
        "version":                version_match.group(1) if version_match else "",
        "author":                 "",
        "date":                   date_match.group(0) if date_match else "",
        "department":             "",
        "project_code":           "",
        "status":                 "",
        "region":                 "",
        "country":                "",
        "geography":              "",
        "language":               "",
        "product_owner":          "",
        "sme":                    "",
        "contributors":           [],
        "elevate_last_edit_date": "",
        "elevate_fields_changed": "",
        "reviewers":              [],
        "format":                 format,
    }

    kv_pattern = re.compile(r"^\*?\*?([A-Za-z ]{3,30})\*?\*?\s*[:\|]\s*(.+)$")
    kv_map = {
        "product owner": "product_owner", "sme": "sme", "status": "status",
        "region": "region", "country": "country", "geography": "geography",
        "language": "language", "version": "version", "author": "author",
        "contributors": "contributors", "content category": "department",
        "elevate last edit date": "elevate_last_edit_date",
        "elevate fields changed": "elevate_fields_changed",
    }
    for line in lines:
        m = kv_pattern.match(line)
        if m:
            key = m.group(1).lower().strip()
            val = m.group(2).strip()
            for pattern, field in kv_map.items():
                if pattern in key:
                    if field == "contributors":
                        metadata["contributors"] = [
                            v.strip() for v in re.split(r"[\n,]+", val) if v.strip()
                        ]
                    else:
                        metadata[field] = val
                    break

    # ── CITATIONS ─────────────────────────────────────────────────────────────
    all_urls = list(dict.fromkeys(url_pattern.findall(joined)))
    references = [
        {
            "id":     f"ref_{i}",
            "type":   "regulation" if "legislation.gov" in url else "document",
            "title":  url,
            "source": url,
        }
        for i, url in enumerate(all_urls[:40], 1)
    ]
    citations: dict = {
        "citation_style":      "Hierarchical pipe-separated citation format (Level2 | Level3 | ...)",
        "references":          references,
        "internal_references": [u for u in all_urls if "confluence" in u.lower() or "file:///" in u.lower()],
        "external_standards":  [u for u in all_urls if "legislation.gov.au" in u],
    }

    # ── CONTENT PROFILE ───────────────────────────────────────────────────────
    req_count = len(re.findall(r"\b(must|shall|required|mandatory|should)\b", lowered))
    theme_patterns = {
        "legislation":   r"legislat|regulation|act\b",
        "metadata":      r"metadata|author|publication date|version",
        "citations":     r"citation|reference",
        "compliance":    r"compliance|regulatory|audit",
        "structuring":   r"table of contents|toc|heading",
        "file_delivery": r"file naming|zip|delivery",
        "australia":     r"australia|au\.|apac",
    }
    key_themes   = [t for t, p in theme_patterns.items() if re.search(p, lowered)]
    complexity   = "high" if len(lines) > 800 else "medium" if len(lines) > 300 else "low"
    expected     = {"scope", "metadata", "citation", "file", "structur", "exception", "update"}
    completeness = round(sum(1 for s in expected if s in lowered) / len(expected) * 100)

    content_profile: dict = {
        "document_type":      "BRD",
        "complexity":         complexity,
        "primary_domain":     (
            "Regulatory / Legal (Australian Legislative Instruments)"
            if "legislation" in lowered else "Business"
        ),
        "key_themes":         key_themes,
        "functional_areas":   [
            l for l in lines if len(l) < 120 and not url_pattern.search(l)
        ][:10],
        "requirements_count": req_count,
        "has_diagrams":       False,
        "has_tables":         "|" in joined or "\t" in joined,
        "completeness_score": completeness,
        "quality_notes":      [],
        "word_count":         len(joined.split()),
    }

    language = metadata.get("language") or "English"

    scope_dict: dict = {
        "in_scope":     in_scope,
        "out_of_scope": [],
        "summary":      f"Scope covers {len(in_scope)} active documents.",
    }

    return {
        "extracted_at":    datetime.now().isoformat() + "Z",
        "source_file":     "",
        "toc":             {"sections": toc_sections},
        "metadata":        metadata,
        "scope":           scope_dict,
        "citations":       citations,
        "content_profile": content_profile,
        "brd_config":      None,
        "cell_images":     [],
    }